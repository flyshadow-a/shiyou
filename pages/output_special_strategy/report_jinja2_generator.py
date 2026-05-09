from __future__ import annotations

import argparse
import copy
import io
import json
import math
import re
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from statistics import NormalDist
from typing import Any

from jinja2 import Environment, StrictUndefined
from openpyxl import load_workbook
import pythoncom
import win32com.client
import xml.etree.ElementTree as ET


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
HYPERLINK_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
STD_NORM = NormalDist()
WORD_NAMESPACE_PREFIXES = {
    "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "cx": "http://schemas.microsoft.com/office/drawing/2014/chartex",
    "cx1": "http://schemas.microsoft.com/office/drawing/2015/9/8/chartex",
    "cx2": "http://schemas.microsoft.com/office/drawing/2015/10/21/chartex",
    "cx3": "http://schemas.microsoft.com/office/drawing/2016/5/9/chartex",
    "cx4": "http://schemas.microsoft.com/office/drawing/2016/5/10/chartex",
    "cx5": "http://schemas.microsoft.com/office/drawing/2016/5/11/chartex",
    "cx6": "http://schemas.microsoft.com/office/drawing/2016/5/12/chartex",
    "cx7": "http://schemas.microsoft.com/office/drawing/2016/5/13/chartex",
    "cx8": "http://schemas.microsoft.com/office/drawing/2016/5/14/chartex",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "aink": "http://schemas.microsoft.com/office/drawing/2016/ink",
    "am3d": "http://schemas.microsoft.com/office/drawing/2017/model3d",
    "o": "urn:schemas-microsoft-com:office:office",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "w10": "urn:schemas-microsoft-com:office:word",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16du": "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16sdtfl": "http://schemas.microsoft.com/office/word/2024/wordml/sdtformatlock",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi": "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "sl": "http://schemas.openxmlformats.org/schemaLibrary/2006/main",
}


@dataclass(frozen=True)
class TableLocator:
    title_keywords: tuple[str, ...]
    header_keywords: tuple[str, ...] = ()
    fallback_index: int | None = None


@dataclass(frozen=True)
class DynamicTableSpec:
    locator: TableLocator
    header_rows: int
    context_key: str
    column_templates: list[str]


DYNAMIC_TABLE_SPECS: list[DynamicTableSpec] = [
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u8282\u70b9\u710a\u7f1d\u75b2\u52b3\u5931\u6548\u6982\u7387",),
            header_keywords=("JointID", "Brace", "JointType"),
            fallback_index=3,
        ),
        header_rows=2,
        context_key="fatigue_failure_rows",
        column_templates=[
            "{{ row.joint_id }}",
            "{{ row.brace }}",
            "{{ row.joint_type }}",
            "{{ row.damage }}",
            "{{ row.beta }}",
            "{{ row.pf }}",
            "{{ row.fatigue_prob_level }}",
            "{{ row.time_node }}",
        ],
    ),
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u5012\u584c\u5206\u6790\u6784\u4ef6\u5931\u6548\u6982\u7387",),
            header_keywords=("Joint", "\u5931\u6548\u6982\u7387"),
            fallback_index=4,
        ),
        header_rows=3,
        context_key="collapse_member_rows",
        column_templates=[
            "{{ row.joint_a }}",
            "{{ row.joint_b }}",
            "{{ row.member_type }}",
            "{{ row.a }}",
            "{{ row.b }}",
            "{{ row.rm }}",
            "{{ row.vr }}",
            "{{ row.pf }}",
            "{{ row.collapse_prob_level }}",
        ],
    ),
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u5012\u584c\u5206\u6790\u8282\u70b9\u5931\u6548\u6982\u7387",),
            header_keywords=("\u8282\u70b9", "\u5931\u6548\u6982\u7387"),
            fallback_index=5,
        ),
        header_rows=2,
        context_key="collapse_joint_rows",
        column_templates=[
            "{{ row.joint_id }}",
            "{{ row.joint_type }}",
            "{{ row.a }}",
            "{{ row.b }}",
            "{{ row.rm }}",
            "{{ row.vr }}",
            "{{ row.pf }}",
            "{{ row.collapse_prob_level }}",
        ],
    ),
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u8282\u70b9\u98ce\u9669\u7b49\u7ea7",),
            header_keywords=("JointID", "Brace", "JointType"),
            fallback_index=7,
        ),
        header_rows=1,
        context_key="node_risk_rows_current",
        column_templates=[
            "{{ row.joint_id }}",
            "{{ row.brace }}",
            "{{ row.joint_type }}",
            "{{ row.consequence_level }}",
            "{{ row.collapse_prob_level }}",
            "{{ row.fatigue_prob_level }}",
            "{{ row.combined_prob_level }}",
            "{{ row.node_risk_level }}",
            "{{ row.time_node }}",
        ],
    ),
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u6784\u4ef6\u98ce\u9669\u7b49\u7ea7",),
            header_keywords=("\u6784\u4ef6", "\u6784\u4ef6\u98ce\u9669\u7b49\u7ea7"),
            fallback_index=8,
        ),
        header_rows=2,
        context_key="member_risk_rows",
        column_templates=[
            "{{ row.joint_a }}",
            "{{ row.joint_b }}",
            "{{ row.member_type }}",
            "{{ row.consequence_level }}",
            "{{ row.collapse_prob_level }}",
            "{{ row.member_risk_level }}",
        ],
    ),
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u6784\u4ef6\u68c0\u6d4b\u8ba1\u5212",),
            header_keywords=("JointA", "JointB", "MemberType"),
            fallback_index=15,
        ),
        header_rows=1,
        context_key="member_inspection_rows",
        column_templates=[
            "{{ row.joint_a }}",
            "{{ row.joint_b }}",
            "{{ row.member_type }}",
            "{{ row.consequence_level }}",
            "{{ row.member_risk_level }}",
            "{{ row.inspect_level }}",
            "{{ row.time_node }}",
        ],
    ),
    DynamicTableSpec(
        locator=TableLocator(
            title_keywords=("\u8282\u70b9\u68c0\u9a8c\u8ba1\u5212",),
            header_keywords=("JointID", "Brace", "JointType"),
            fallback_index=16,
        ),
        header_rows=1,
        context_key="node_inspection_rows_future",
        column_templates=[
            "{{ row.joint_id }}",
            "{{ row.brace }}",
            "{{ row.joint_type }}",
            "{{ row.consequence_level }}",
            "{{ row.combined_prob_level }}",
            "{{ row.node_risk_level }}",
            "{{ row.inspect_level }}",
            "{{ row.time_node }}",
        ],
    ),
]


SUMMARY_TABLE_LOCATORS: dict[str, TableLocator] = {
    "member_risk_summary": TableLocator(
        title_keywords=("\u5e73\u53f0\u6784\u4ef6\u98ce\u9669\u7b49\u7ea7\u6c47\u603b",),
        fallback_index=9,
    ),
    "node_risk_summary": TableLocator(
        title_keywords=("\u5e73\u53f0\u8282\u70b9\u98ce\u9669\u7b49\u7ea7\u6c47\u603b",),
        fallback_index=10,
    ),
    "member_inspection_summary": TableLocator(
        title_keywords=("\u6784\u4ef6\u68c0\u9a8c\u8ba1\u5212\u6c47\u603b\u8868",),
        fallback_index=13,
    ),
    "node_inspection_summary": TableLocator(
        title_keywords=("\u8282\u70b9\u68c0\u9a8c\u8ba1\u5212\u6c47\u603b\u8868",),
        fallback_index=14,
    ),
}


RISK_LEVEL_ORDER = ["\u4e00", "\u4e8c", "\u4e09", "\u56db", "\u4e94"]
INSPECTION_LEVEL_ORDER = ["II", "III", "IV"]
TIME_CURRENT = "\u5f53\u524d"
TIME_ORDER = [TIME_CURRENT, "\u7b2c5\u5e74", "\u7b2c10\u5e74", "\u7b2c15\u5e74", "\u7b2c20\u5e74", "\u7b2c25\u5e74"]
TIME_TO_YEAR = {
    TIME_ORDER[1]: 5,
    TIME_ORDER[2]: 10,
    TIME_ORDER[3]: 15,
    TIME_ORDER[4]: 20,
    TIME_ORDER[5]: 25,
}

DEFAULT_DYNAMIC_ROW_LIMITS = {
    "fatigue_failure_rows": 800,
    "collapse_member_rows": 80,
    "collapse_joint_rows": 80,
    "node_risk_rows_current": 250,
    "member_risk_rows": 250,
    "member_inspection_rows": 250,
    "node_inspection_rows_future": 500,
}

OPTIONAL_WORD_DETAIL_CONTEXT_KEYS = {
    "member_inspection_rows",
    "node_inspection_rows_future",
}


APPENDIX_A_HEADING = "附件A 节点风险等级计算表"
APPENDIX_B_HEADING = "附件B 构件风险等级计算表"
APPENDIX_C_MAIN_HEADING = "附件C 平台检验计划的检验节点构件"
APPENDIX_C_SUBHEADINGS = [
    "附件C.1 服役第5年的检验节点构件位置",
    "附件C.2 服役第10年的检验节点构件位置",
    "附件C.3 服役第15年的检验节点构件位置",
    "附件C.4 服役第20年的检验节点构件位置",
    "附件C.5 服役第25年的检验节点构件位置",
]
APPENDIX_IMAGE_WIDTH_INCHES = 6.2
# Lower appendix render quality to keep the final DOCX responsive when many
# PDF pages are embedded. The report body stays unchanged; only appendix image
# density/compression is reduced.
APPENDIX_RENDER_DPI = 90
APPENDIX_JPEG_QUALITY = 60
GENERATED_APPENDIX_SHEETS = [
    (APPENDIX_A_HEADING, "节点失效风险等级", "appendix_a_node_risk.pdf"),
    (APPENDIX_B_HEADING, "构件失效风险等级", "appendix_b_member_risk.pdf"),
]
APPENDIX_C_YEAR_SECTION_MAP = [
    ("+5年", APPENDIX_C_SUBHEADINGS[0]),
    ("+10年", APPENDIX_C_SUBHEADINGS[1]),
    ("+15年", APPENDIX_C_SUBHEADINGS[2]),
    ("+20年", APPENDIX_C_SUBHEADINGS[3]),
    ("+25年", APPENDIX_C_SUBHEADINGS[4]),
]
APPENDIX_C_ROW_NAMES = ["XZ 前", "XZ 后", "YZ 左", "YZ 右"]
APPENDIX_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}


def to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value).strip()


def register_word_namespaces() -> None:
    ET.register_namespace("", REL_NS)
    for prefix, uri in WORD_NAMESPACE_PREFIXES.items():
        ET.register_namespace(prefix, uri)


def to_scientific_text(value: Any, digits: int = 2) -> str:
    """
    Format numeric value in scientific notation (e.g. 7.03E-42).
    Keep empty/non-numeric values unchanged.
    """
    if value is None:
        return ""
    txt = str(value).strip()
    if txt == "":
        return ""
    try:
        f = float(txt)
    except ValueError:
        return txt
    if math.isnan(f):
        return ""
    return f"{f:.{digits}E}"


def to_int_or_none(value: Any) -> int | None:
    txt = to_text(value)
    if txt == "":
        return None
    try:
        return int(float(txt))
    except ValueError:
        return None


def to_float_or_none(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        f = float(value)
        if math.isnan(f):
            return None
        return f
    txt = str(value).strip()
    if txt == "":
        return None
    try:
        f = float(txt)
    except ValueError:
        return None
    if math.isnan(f):
        return None
    return f


def collapse_pf_from_factor(a: Any, b: Any, rm: Any, vr: Any) -> float | None:
    a_f = to_float_or_none(a)
    b_f = to_float_or_none(b)
    rm_f = to_float_or_none(rm)
    vr_f = to_float_or_none(vr)
    if a_f is None or b_f is None or rm_f is None or vr_f is None:
        return None
    if b_f == 0.0 or vr_f == 0.0:
        return None
    term1 = math.exp((a_f - rm_f) / b_f + (vr_f**2) * (rm_f**2) / (2 * (b_f**2)))
    # VBA uses NORM.DIST(..., FALSE): standard normal PDF.
    term2 = 1.0 - STD_NORM.pdf(vr_f * rm_f / b_f - 1.0 / vr_f)
    return float(term1 * term2)


def normalize_time_node(raw: Any) -> str:
    s = to_text(raw)
    if not s:
        return ""
    nums = re.findall(r"\d+", s)
    if not nums:
        return TIME_CURRENT
    return f"\u7b2c{int(nums[0])}\u5e74"


PLAN_TIME_MAP = {
    "N": TIME_CURRENT,
    "N+5": "\u7b2c5\u5e74",
    "N+10": "\u7b2c10\u5e74",
    "N+15": "\u7b2c15\u5e74",
    "N+20": "\u7b2c20\u5e74",
    "N+25": "\u7b2c25\u5e74",
}

FORECAST_YEAR_TIME_MAP = {
    1: TIME_CURRENT,
    6: "\u7b2c5\u5e74",
    11: "\u7b2c10\u5e74",
    16: "\u7b2c15\u5e74",
    21: "\u7b2c20\u5e74",
    26: "\u7b2c25\u5e74",
}

RISK_LEVEL_TO_INDEX = {lv: idx for idx, lv in enumerate(RISK_LEVEL_ORDER)}


def normalize_plan_time(raw: Any) -> str:
    s = to_text(raw).upper().replace(" ", "")
    if s in PLAN_TIME_MAP:
        return PLAN_TIME_MAP[s]
    return normalize_time_node(s)


def forecast_year_to_time_node(year: int | None) -> str:
    if year is None:
        return ""
    if year in FORECAST_YEAR_TIME_MAP:
        return FORECAST_YEAR_TIME_MAP[year]
    if year <= 1:
        return TIME_CURRENT
    return f"\u7b2c{year - 1}\u5e74"


def normalize_inspect_level(raw: Any) -> str:
    level = to_text(raw).upper()
    if level in {"II", "III", "IV"}:
        return level
    if level == "I":
        return "II"
    return "III"


def choose_higher_risk_grade(grade_a: Any, grade_b: Any) -> str:
    a = to_text(grade_a)
    b = to_text(grade_b)
    candidates = [g for g in (a, b) if g in RISK_LEVEL_TO_INDEX]
    if not candidates:
        return ""
    return min(candidates, key=lambda g: RISK_LEVEL_TO_INDEX[g])


def collapse_level_to_risk_grade(level: int | None) -> str:
    if level is None:
        return ""
    if 1 <= level <= len(RISK_LEVEL_ORDER):
        return RISK_LEVEL_ORDER[level - 1]
    return ""


def risk_rank(risk: Any) -> int:
    s = to_text(risk)
    if s in RISK_LEVEL_TO_INDEX:
        return RISK_LEVEL_TO_INDEX[s]
    return len(RISK_LEVEL_ORDER) + 1


def sorted_node_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(
        rows,
        key=lambda r: (
            risk_rank(r.get("node_risk_level")),
            to_int_or_none(r.get("combined_prob_level")) or 99,
            to_int_or_none(r.get("fatigue_prob_level")) or 99,
            to_text(r.get("joint_id")),
        ),
    )


def sorted_fatigue_failure_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Table 2-5 sorting:
    1) time node ascending (第5年 -> 第10年 -> 第15年 -> 第20年 -> 第25年)
    2) joint/brace lexicographic
    3) fatigue probability level ascending
    4) node risk level ascending
    """
    return sorted(
        rows,
        key=lambda r: (
            time_node_rank(r.get("time_node")),
            to_text(r.get("joint_id")),
            to_text(r.get("brace")),
            to_int_or_none(r.get("fatigue_prob_level")) or 99,
            risk_rank(r.get("node_risk_level")),
        ),
    )


def sorted_member_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(
        rows,
        key=lambda r: (
            risk_rank(r.get("member_risk_level")),
            to_int_or_none(r.get("collapse_prob_level")) or 99,
            to_text(r.get("joint_a")),
            to_text(r.get("joint_b")),
        ),
    )


def cap_rows(rows: list[dict[str, Any]], limit: int | None) -> list[dict[str, Any]]:
    if limit is None or limit <= 0 or len(rows) <= limit:
        return rows
    return rows[:limit]


def iter_data_rows(ws, min_row: int, max_col: int, key_col_index: int = 0):
    blank_run = 0
    for row in ws.iter_rows(min_row=min_row, max_col=max_col, values_only=True):
        key = row[key_col_index]
        if to_text(key) == "":
            blank_run += 1
            if blank_run > 50:
                break
            continue
        blank_run = 0
        yield row


def get_sheet_by_names_or_index(wb, names: list[str], fallback_index: int):
    for name in names:
        if name in wb.sheetnames:
            return wb[name]
    if 0 <= fallback_index < len(wb.worksheets):
        return wb.worksheets[fallback_index]
    raise KeyError(f"Sheet not found by names={names} or index={fallback_index}")


def build_collapse_overrides_from_sheet(
    ws_collapse,
    ws_members,
    node_type_by_joint: dict[str, str],
    member_level_by_pair: dict[str, str] | None,
    joint_level_by_id: dict[str, str] | None,
    collapse_a_const: Any,
    collapse_b_const: Any,
    collapse_vr_const: Any,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    member_type_by_pair: dict[str, str] = {}
    if ws_members is not None:
        for row in iter_data_rows(ws_members, min_row=2, max_col=5):
            ja = to_text(row[0])
            jb = to_text(row[1])
            mt = to_text(row[4]) or "Other"
            if ja == "" or jb == "":
                continue
            member_type_by_pair[f"{ja}-{jb}"] = mt
            member_type_by_pair[f"{jb}-{ja}"] = mt

    collapse_member_factor: dict[str, float | None] = {}
    collapse_joint_factor: dict[str, float | None] = {}
    if ws_collapse is not None:
        for row in iter_data_rows(ws_collapse, min_row=2, max_col=5):
            fail_type = to_text(row[1])
            location = to_text(row[2])
            factor = to_float_or_none(row[3])
            if location == "":
                continue
            if fail_type.upper() == "TYPE" or location.upper() == "LOCATION":
                continue
            ft = fail_type.replace(" ", "").upper()
            is_member = ("MEMBER" in ft) or ("构件" in fail_type and "失效" in fail_type)
            is_joint = ("JOINT" in ft) or ("节点" in fail_type and "失效" in fail_type)
            if not is_member and not is_joint:
                is_member = "-" in location
                is_joint = not is_member

            if is_member:
                old = collapse_member_factor.get(location)
                if old is None or (factor is not None and factor < old):
                    collapse_member_factor[location] = factor
            elif is_joint:
                old = collapse_joint_factor.get(location)
                if old is None or (factor is not None and factor < old):
                    collapse_joint_factor[location] = factor

    collapse_member_rows_override: list[dict[str, Any]] = []
    for location, factor in sorted(
        collapse_member_factor.items(),
        key=lambda kv: (
            9999 if kv[1] is None else kv[1],
            kv[0],
        ),
    ):
        ja = ""
        jb = ""
        if "-" in location:
            ja, jb = location.split("-", 1)
        mtype = member_type_by_pair.get(location) or "Other"
        pf_val = collapse_pf_from_factor(collapse_a_const, collapse_b_const, factor, collapse_vr_const)
        level = ""
        if member_level_by_pair:
            level = to_text(member_level_by_pair.get(location))
            if level == "" and ja != "" and jb != "":
                level = to_text(member_level_by_pair.get(f"{jb}-{ja}"))

        a_txt = to_text(collapse_a_const)
        b_txt = to_text(collapse_b_const)
        a_val = to_float_or_none(collapse_a_const)
        b_val = to_float_or_none(collapse_b_const)
        if a_val is not None:
            a_txt = f"{a_val:.3f}"
        if b_val is not None:
            b_txt = f"{b_val:.3f}"

        vr_txt = to_text(collapse_vr_const)
        vr_val = to_float_or_none(collapse_vr_const)
        if vr_val is not None:
            vr_txt = f"{vr_val * 100:.0f}%"

        collapse_member_rows_override.append(
            {
                "joint_a": ja,
                "joint_b": jb,
                "member_type": mtype,
                "a": a_txt,
                "b": b_txt,
                "rm": "" if factor is None else f"{factor:g}",
                "vr": vr_txt,
                "pf": to_scientific_text(pf_val),
                "collapse_prob_level": level,
            }
        )

    collapse_joint_rows_override: list[dict[str, Any]] = []
    for joint_id, factor in sorted(
        collapse_joint_factor.items(),
        key=lambda kv: (
            9999 if kv[1] is None else kv[1],
            kv[0],
        ),
    ):
        pf_val = collapse_pf_from_factor(collapse_a_const, collapse_b_const, factor, collapse_vr_const)
        level = ""
        if joint_level_by_id:
            level = to_text(joint_level_by_id.get(joint_id))
        lv_i = to_int_or_none(level)
        if lv_i is not None and lv_i > 4:
            continue

        a_txt = to_text(collapse_a_const)
        b_txt = to_text(collapse_b_const)
        a_val = to_float_or_none(collapse_a_const)
        b_val = to_float_or_none(collapse_b_const)
        if a_val is not None:
            a_txt = f"{a_val:.3f}"
        if b_val is not None:
            b_txt = f"{b_val:.3f}"

        vr_txt = to_text(collapse_vr_const)
        vr_val = to_float_or_none(collapse_vr_const)
        if vr_val is not None:
            vr_txt = f"{vr_val * 100:.0f}%"

        collapse_joint_rows_override.append(
            {
                "joint_id": joint_id,
                "joint_type": node_type_by_joint.get(joint_id, ""),
                "a": a_txt,
                "b": b_txt,
                "rm": "" if factor is None else f"{factor:g}",
                "vr": vr_txt,
                "pf": to_scientific_text(pf_val),
                "collapse_prob_level": level,
            }
        )

    return collapse_member_rows_override, collapse_joint_rows_override


def aggregate_risk_counts(rows: list[dict[str, Any]], risk_key: str) -> dict[str, int]:
    counts = {k: 0 for k in RISK_LEVEL_ORDER}
    for row in rows:
        risk = to_text(row.get(risk_key))
        if risk in counts:
            counts[risk] += 1
    return counts


def aggregate_inspection_counts(rows: list[dict[str, Any]], risk_value: str, risk_key: str) -> dict[str, int]:
    counts = {k: 0 for k in INSPECTION_LEVEL_ORDER}
    for row in rows:
        if to_text(row.get(risk_key)) != risk_value:
            continue
        level = to_text(row.get("inspect_level"))
        if level in counts:
            counts[level] += 1
    return counts


def time_node_rank(value: Any) -> int:
    """
    Rank time node for earliest-first selection.
    """
    s = normalize_time_node(value)
    if s in TIME_ORDER:
        return TIME_ORDER.index(s)
    nums = re.findall(r"\d+", to_text(value))
    if nums:
        return int(nums[0])
    return 9999


def is_deleted_joint_by_vba_rule(joint_id: Any) -> bool:
    """
    VBA Sheet11/Sheet12 删除JOINT* rules:
    delete when A matches one of:
    - C*
    - B*
    - J*
    - O*
    - K###
    - R*
    """
    s = to_text(joint_id).upper()
    if s == "":
        return False
    if s.startswith(("C", "B", "J", "O", "R")):
        return True
    if re.fullmatch(r"K\d{3}", s):
        return True
    return False


def is_deleted_member_by_vba_rule(joint_a: Any, joint_b: Any) -> bool:
    """
    VBA Sheet10 删除MEMBER rules:
    delete row when any condition matches:
    - A Like "C*" Or B Like "C*"
    - A Like "B*" And B Like "B*"
    - A Like "J*" Or B Like "J*"
    - A Like "O*" Or B Like "O*"
    - A Like "K###" And B Like "K###"
    - A Like "R*" Or B Like "R*"
    - A Like "0*L" And B Like "0*L"
    """
    a = to_text(joint_a).upper()
    b = to_text(joint_b).upper()
    if a == "" or b == "":
        return False
    if a.startswith("C") or b.startswith("C"):
        return True
    if a.startswith("B") and b.startswith("B"):
        return True
    if a.startswith("J") or b.startswith("J"):
        return True
    if a.startswith("O") or b.startswith("O"):
        return True
    if re.fullmatch(r"K\d{3}", a) and re.fullmatch(r"K\d{3}", b):
        return True
    if a.startswith("R") or b.startswith("R"):
        return True
    if a.startswith("0") and a.endswith("L") and b.startswith("0") and b.endswith("L"):
        return True
    return False


def ratio_text(part: int, total: int) -> str:
    if total <= 0:
        return "0.00%"
    return f"{(part / total) * 100:.2f}%"


def merge_metadata_into_context(context: dict[str, Any], metadata: dict[str, Any] | None) -> dict[str, Any]:
    payload = dict(context)
    metadata_payload = dict(metadata or {})
    payload["report_metadata"] = metadata_payload

    for key, value in metadata_payload.items():
        if value in ("", None):
            continue
        if key in ("platform_name", "report_date") or key not in payload:
            payload[str(key)] = value
    return payload


def build_context(
    node_risk_rows: list[dict[str, Any]],
    node_strategy_rows: list[dict[str, Any]],
    member_risk_rows: list[dict[str, Any]],
    member_strategy_rows: list[dict[str, Any]],
    metadata: dict[str, Any],
    row_limits: dict[str, int] | None = None,
    collapse_member_rows_override: list[dict[str, Any]] | None = None,
    collapse_joint_rows_override: list[dict[str, Any]] | None = None,
    apply_vba_delete_rules: bool = False,
    apply_vba_member_delete_rules: bool | None = None,
    apply_vba_joint_delete_rules: bool | None = None,
    apply_vba_joint_delete_rules_current: bool | None = None,
    apply_vba_joint_delete_rules_future: bool | None = None,
) -> dict[str, Any]:
    if row_limits is None:
        limits = {k: None for k in DEFAULT_DYNAMIC_ROW_LIMITS}
    else:
        limits = dict(DEFAULT_DYNAMIC_ROW_LIMITS)
        limits.update(row_limits)

    if apply_vba_member_delete_rules is None:
        apply_vba_member_delete_rules = apply_vba_delete_rules
    if apply_vba_joint_delete_rules is None:
        apply_vba_joint_delete_rules = apply_vba_delete_rules
    if apply_vba_joint_delete_rules_current is None:
        apply_vba_joint_delete_rules_current = apply_vba_joint_delete_rules
    if apply_vba_joint_delete_rules_future is None:
        apply_vba_joint_delete_rules_future = apply_vba_joint_delete_rules

    if apply_vba_joint_delete_rules_current:
        node_risk_rows_use = [r for r in node_risk_rows if not is_deleted_joint_by_vba_rule(r.get("joint_id"))]
    else:
        node_risk_rows_use = list(node_risk_rows)

    if apply_vba_joint_delete_rules_current or apply_vba_joint_delete_rules_future:
        node_strategy_rows_use = []
        for r in node_strategy_rows:
            jid = r.get("joint_id")
            if not is_deleted_joint_by_vba_rule(jid):
                node_strategy_rows_use.append(r)
                continue

            t = normalize_time_node(r.get("time_node"))
            is_current = (t == TIME_CURRENT) or (t == "")
            if is_current:
                if not apply_vba_joint_delete_rules_current:
                    node_strategy_rows_use.append(r)
            else:
                if not apply_vba_joint_delete_rules_future:
                    node_strategy_rows_use.append(r)
    else:
        node_strategy_rows_use = list(node_strategy_rows)
        node_strategy_rows_use = list(node_strategy_rows)

    if apply_vba_member_delete_rules:
        member_risk_rows_use = [
            r
            for r in member_risk_rows
            if not is_deleted_member_by_vba_rule(r.get("joint_a"), r.get("joint_b"))
        ]
        member_strategy_rows_use = [
            r
            for r in member_strategy_rows
            if not is_deleted_member_by_vba_rule(r.get("joint_a"), r.get("joint_b"))
        ]
    else:
        member_risk_rows_use = list(member_risk_rows)
        member_strategy_rows_use = list(member_strategy_rows)

    collapse_member_rows = []
    for r in member_risk_rows_use:
        lv = to_int_or_none(r["collapse_prob_level"])
        if lv is not None and lv <= 2:
            collapse_member_rows.append(r)
    if collapse_member_rows_override is not None:
        collapse_member_rows = collapse_member_rows_override

    collapse_joint_rows = []
    for r in node_risk_rows_use:
        lv = to_int_or_none(r["collapse_prob_level"])
        if lv is not None and lv <= 2:
            collapse_joint_rows.append(r)
    if collapse_joint_rows_override is not None:
        collapse_joint_rows = collapse_joint_rows_override

    # Table 2-5: keep all future rows where fatigue probability level reaches
    # failure threshold (<=2). Keep all joint types, including "Other".
    # Do not deduplicate by (JointID, Brace); keep per-time-node rows.
    fatigue_failure_rows: list[dict[str, Any]] = []
    for r in node_strategy_rows_use:
        if r["time_node"] == TIME_CURRENT:
            continue
        lv = to_int_or_none(r.get("fatigue_prob_level"))
        if lv is not None and lv <= 2:
            fatigue_failure_rows.append(r)

    if not fatigue_failure_rows:
        fatigue_failure_rows = [r for r in node_strategy_rows_use if r["time_node"] != TIME_CURRENT]

    node_risk_rows_current = [r for r in node_strategy_rows_use if r["time_node"] == TIME_CURRENT]
    # Word Table 4-4 in the reference report only lists future checkpoints that
    # require elevated inspection methods, instead of the full node strategy table.
    node_inspection_rows_future = [
        r
        for r in node_strategy_rows_use
        if r["time_node"] != TIME_CURRENT and to_text(r.get("inspect_level")) in ("III", "IV")
    ]
    # Table 2-9 (表29): future nodes with risk level 一/二.
    # Keep source sheet order for strict VBA-style output.
    node_risk_rows_table29 = [
        r for r in node_strategy_rows_use
        if r["time_node"] != TIME_CURRENT and risk_rank(r.get("node_risk_level")) <= 1
    ]
    # Fallback to combined probability <=3 (equivalent to high-risk zone in most matrices).
    if not node_risk_rows_table29:
        node_risk_rows_table29 = [
            r for r in node_strategy_rows_use
            if r["time_node"] != TIME_CURRENT and (to_int_or_none(r.get("combined_prob_level")) or 99) <= 3
        ]
    if not node_risk_rows_table29:
        node_risk_rows_table29 = [r for r in node_strategy_rows_use if r["time_node"] != TIME_CURRENT]

    member_risk_counts = aggregate_risk_counts(member_risk_rows_use, "member_risk_level")
    total_member = sum(member_risk_counts.values())
    member_risk_ratios = {k: ratio_text(v, total_member) for k, v in member_risk_counts.items()}

    node_summary_blocks = []
    for t in TIME_ORDER:
        rows_t = [r for r in node_strategy_rows_use if r["time_node"] == t]
        if not rows_t:
            continue
        counts = aggregate_risk_counts(rows_t, "node_risk_level")
        total = sum(counts.values())
        ratios = {k: ratio_text(v, total) for k, v in counts.items()}
        if t == TIME_CURRENT:
            title = "\u8282\u70b9\u710a\u7f1d\uff08\u5f53\u524d\uff09"
        else:
            year = TIME_TO_YEAR.get(t, "")
            title = f"\u8282\u70b9\u710a\u7f1d\uff08\u672a\u6765{year}\u5e74\uff09"
        node_summary_blocks.append({"time_node": t, "title": title, "counts": counts, "ratios": ratios, "total": total})

    # Word Table 4-1 summarizes the current member plan only. Future-year member
    # strategy rows are shown separately in Table 4-3.
    member_strategy_rows_current = [
        r for r in member_strategy_rows_use if normalize_time_node(r.get("time_node")) in (TIME_CURRENT, "")
    ]
    member_inspection_summary = []
    member_total = 0
    member_level_totals = {k: 0 for k in INSPECTION_LEVEL_ORDER}
    for risk in RISK_LEVEL_ORDER:
        rows_r = [r for r in member_strategy_rows_current if to_text(r["member_risk_level"]) == risk]
        risk_count = len(rows_r)
        inspect_counts = aggregate_inspection_counts(rows_r, risk, "member_risk_level")
        member_total += risk_count
        for lv in INSPECTION_LEVEL_ORDER:
            member_level_totals[lv] += inspect_counts[lv]
        member_inspection_summary.append(
            {
                "risk_level": risk,
                "count": risk_count,
                "II": inspect_counts["II"],
                "III": inspect_counts["III"],
                "IV": inspect_counts["IV"],
            }
        )

    node_inspection_blocks = []
    for t in [x for x in TIME_ORDER if x != TIME_CURRENT]:
        rows_t = [r for r in node_strategy_rows_use if r["time_node"] == t]
        if not rows_t:
            continue
        summary_rows = []
        total = len(rows_t)
        total_level_counts = {k: 0 for k in INSPECTION_LEVEL_ORDER}
        for risk in RISK_LEVEL_ORDER:
            rows_r = [r for r in rows_t if to_text(r["node_risk_level"]) == risk]
            inspect_counts = aggregate_inspection_counts(rows_r, risk, "node_risk_level")
            for lv in INSPECTION_LEVEL_ORDER:
                total_level_counts[lv] += inspect_counts[lv]
            summary_rows.append(
                {
                    "risk_level": risk,
                    "count": len(rows_r),
                    "II": inspect_counts["II"],
                    "III": inspect_counts["III"],
                    "IV": inspect_counts["IV"],
                }
            )
        node_inspection_blocks.append(
            {
                "time_node": t,
                "summary_rows": summary_rows,
                "total_count": total,
                "total_II": total_level_counts["II"],
                "total_III": total_level_counts["III"],
                "total_IV": total_level_counts["IV"],
            }
        )

    # Keep source-table order for strict VBA alignment; only apply row caps.
    collapse_member_rows_disp = cap_rows(collapse_member_rows, limits.get("collapse_member_rows"))
    collapse_joint_rows_disp = cap_rows(collapse_joint_rows, limits.get("collapse_joint_rows"))
    fatigue_failure_rows_disp = cap_rows(
        sorted_fatigue_failure_rows(fatigue_failure_rows),
        limits.get("fatigue_failure_rows"),
    )
    table29_limit = limits.get("node_risk_rows_current")
    if table29_limit is None or table29_limit <= 0 or len(node_risk_rows_table29) <= table29_limit:
        node_risk_rows_current_disp = list(node_risk_rows_table29)
    else:
        # Keep every future checkpoint visible under row cap:
        # distribute slots across 5/10/15/20/25 years.
        table29_groups = {
            t: [r for r in node_risk_rows_table29 if to_text(r.get("time_node")) == t]
            for t in TIME_ORDER
            if t != TIME_CURRENT
        }
        active_times = [t for t in TIME_ORDER if t != TIME_CURRENT and len(table29_groups.get(t, [])) > 0]
        if not active_times:
            node_risk_rows_current_disp = cap_rows(node_risk_rows_table29, table29_limit)
        else:
            per_time = max(1, table29_limit // len(active_times))
            picked: dict[str, int] = {}
            used = 0
            for t in active_times:
                take = min(len(table29_groups[t]), per_time)
                picked[t] = take
                used += take
            remain = max(0, table29_limit - used)
            while remain > 0:
                progressed = False
                for t in active_times:
                    if picked[t] < len(table29_groups[t]):
                        picked[t] += 1
                        remain -= 1
                        progressed = True
                        if remain == 0:
                            break
                if not progressed:
                    break
            node_risk_rows_current_disp = []
            for t in active_times:
                node_risk_rows_current_disp.extend(table29_groups[t][: picked[t]])

    # Word Table 4-3 shows only future member rows that need III/IV inspection.
    member_inspection_rows = [
        r
        for r in member_strategy_rows_use
        if normalize_time_node(r.get("time_node")) not in (TIME_CURRENT, "")
        and to_text(r.get("inspect_level")) in ("III", "IV")
    ]

    member_risk_rows_disp = cap_rows(member_risk_rows_use, limits.get("member_risk_rows"))
    member_inspection_rows_disp = cap_rows(member_inspection_rows, limits.get("member_inspection_rows"))
    node_inspection_rows_future_disp = cap_rows(node_inspection_rows_future, limits.get("node_inspection_rows_future"))

    context = {
        "platform_name": metadata.get("platform_name", ""),
        "report_date": metadata.get("report_date", ""),
        "fatigue_failure_rows": fatigue_failure_rows_disp,
        "collapse_member_rows": collapse_member_rows_disp,
        "collapse_joint_rows": collapse_joint_rows_disp,
        "node_risk_rows_current": node_risk_rows_current_disp,
        "member_risk_rows": member_risk_rows_disp,
        "member_inspection_rows": member_inspection_rows_disp,
        "node_inspection_rows_future": node_inspection_rows_future_disp,
        "member_risk_counts": member_risk_counts,
        "member_risk_ratios": member_risk_ratios,
        "node_summary_blocks": node_summary_blocks,
        "member_inspection_summary": member_inspection_summary,
        "member_inspection_total": member_total,
        "member_inspection_total_II": member_level_totals["II"],
        "member_inspection_total_III": member_level_totals["III"],
        "member_inspection_total_IV": member_level_totals["IV"],
        "node_inspection_blocks": node_inspection_blocks,
        "row_limits": limits,
        "row_counts": {
            "fatigue_failure_rows_total": len(fatigue_failure_rows),
            "collapse_member_rows_total": len(collapse_member_rows),
            "collapse_joint_rows_total": len(collapse_joint_rows),
            "node_risk_rows_current_total": len(node_risk_rows_table29),
            "member_risk_rows_total": len(member_risk_rows_use),
            "member_inspection_rows_total": len(member_inspection_rows),
            "node_inspection_rows_future_total": len(node_inspection_rows_future),
        },
    }
    return merge_metadata_into_context(context, metadata)


def load_context_from_legacy_workbook(
    wb,
    metadata: dict[str, Any],
    row_limits: dict[str, int] | None = None,
    apply_vba_member_delete_rules: bool = False,
    apply_vba_joint_delete_rules_current: bool = False,
    apply_vba_joint_delete_rules_future: bool = False,
) -> dict[str, Any]:
    ws_node_risk = get_sheet_by_names_or_index(wb, ["节点失效风险等级"], 9)
    ws_node_strategy = get_sheet_by_names_or_index(wb, ["节点检验策略"], 11)
    ws_member_risk = get_sheet_by_names_or_index(wb, ["构件失效风险等级"], 12)
    ws_member_strategy = get_sheet_by_names_or_index(wb, ["构件检验策略"], 13)
    ws_collapse = wb["倒塌分析结果"] if "倒塌分析结果" in wb.sheetnames else (wb.worksheets[6] if len(wb.worksheets) > 6 else None)
    ws_members = wb["Members"] if "Members" in wb.sheetnames else (wb.worksheets[17] if len(wb.worksheets) > 17 else None)

    # Re-assign with unicode escapes to avoid source-encoding corruption.
    ws_node_risk = get_sheet_by_names_or_index(wb, ["\u8282\u70b9\u5931\u6548\u98ce\u9669\u7b49\u7ea7"], 9)
    ws_node_strategy = get_sheet_by_names_or_index(wb, ["\u8282\u70b9\u68c0\u9a8c\u7b56\u7565"], 11)
    ws_member_risk = get_sheet_by_names_or_index(wb, ["\u6784\u4ef6\u5931\u6548\u98ce\u9669\u7b49\u7ea7"], 12)
    ws_member_strategy = get_sheet_by_names_or_index(wb, ["\u6784\u4ef6\u68c0\u9a8c\u7b56\u7565"], 13)
    ws_control = wb["\u63a7\u5236\u9875\u9762"] if "\u63a7\u5236\u9875\u9762" in wb.sheetnames else (wb.worksheets[1] if len(wb.worksheets) > 1 else None)
    ws_risk_matrix = wb["\u98ce\u9669\u8bc4\u7ea7\u77e9\u9635"] if "\u98ce\u9669\u8bc4\u7ea7\u77e9\u9635" in wb.sheetnames else (wb.worksheets[2] if len(wb.worksheets) > 2 else None)

    node_risk_rows: list[dict[str, Any]] = []
    node_type_by_joint: dict[str, str] = {}
    for row in iter_data_rows(ws_node_risk, min_row=3, max_col=10):
        joint_id = to_text(row[0])
        one = {
            "joint_id": joint_id,
            "brace": to_text(row[1]),
            "joint_type": to_text(row[2]),
            "consequence_level": to_text(row[3]),
            "a": to_text(row[4]),
            "b": to_text(row[5]),
            "rm": to_text(row[6]),
            "vr": to_text(row[7]),
            "pf": to_scientific_text(row[8]),
            "collapse_prob_level": to_text(row[9]),
        }
        node_risk_rows.append(one)
        if joint_id:
            node_type_by_joint[joint_id] = one["joint_type"]

    node_strategy_rows: list[dict[str, Any]] = []
    for row in iter_data_rows(ws_node_strategy, min_row=3, max_col=18):
        node_strategy_rows.append(
            {
                "joint_id": to_text(row[0]),
                "brace": to_text(row[1]),
                "joint_type": to_text(row[2]),
                "consequence_level": to_text(row[3]),
                "collapse_prob_level": to_text(row[4]),
                "damage": to_scientific_text(row[5]),
                "beta": to_text(row[11]),
                "pf": to_scientific_text(row[12]),
                "fatigue_prob_level": to_text(row[13]),
                "combined_prob_level": to_text(row[14]),
                "node_risk_level": to_text(row[15]),
                "inspect_level": normalize_inspect_level(row[16]),
                "time_node": normalize_time_node(row[17]),
            }
        )

    member_risk_rows: list[dict[str, Any]] = []
    for row in iter_data_rows(ws_member_risk, min_row=3, max_col=11):
        member_risk_rows.append(
            {
                "joint_a": to_text(row[0]),
                "joint_b": to_text(row[1]),
                "member_type": to_text(row[2]),
                "consequence_level": to_text(row[3]),
                "a": to_text(row[4]),
                "b": to_text(row[5]),
                "rm": to_text(row[6]),
                "vr": to_text(row[7]),
                "pf": to_scientific_text(row[8]),
                "collapse_prob_level": to_text(row[9]),
                "member_risk_level": to_text(row[10]),
            }
        )

    member_strategy_rows: list[dict[str, Any]] = []
    for row in iter_data_rows(ws_member_strategy, min_row=2, max_col=7):
        member_strategy_rows.append(
            {
                "joint_a": to_text(row[0]),
                "joint_b": to_text(row[1]),
                "member_type": to_text(row[2]),
                "consequence_level": to_text(row[3]),
                "member_risk_level": to_text(row[4]),
                "inspect_level": normalize_inspect_level(row[5]),
                "time_node": normalize_time_node(row[6]),
            }
        )

    member_level_by_pair: dict[str, str] = {}
    for r in member_risk_rows:
        ja = to_text(r.get("joint_a"))
        jb = to_text(r.get("joint_b"))
        lv = to_text(r.get("collapse_prob_level"))
        if ja != "" and jb != "" and lv != "":
            member_level_by_pair[f"{ja}-{jb}"] = lv
            member_level_by_pair[f"{jb}-{ja}"] = lv

    joint_level_by_id: dict[str, str] = {}
    for r in node_risk_rows:
        jid = to_text(r.get("joint_id"))
        lv = to_text(r.get("collapse_prob_level"))
        if jid != "" and lv != "":
            joint_level_by_id[jid] = lv

    collapse_a_const = None
    collapse_b_const = None
    collapse_vr_const = None
    if ws_control is not None and ws_risk_matrix is not None:
        region_name = to_text(ws_control["B45"].value)
        if region_name != "":
            for c in range(2, 30):
                if to_text(ws_risk_matrix.cell(50, c).value) == region_name:
                    collapse_a_const = ws_risk_matrix.cell(51, c).value
                    collapse_b_const = ws_risk_matrix.cell(52, c).value
                    break
    if ws_control is not None and (collapse_a_const is None or collapse_b_const is None):
        collapse_a_const = ws_control["B46"].value
        collapse_b_const = ws_control["B47"].value
    if member_risk_rows and (collapse_a_const is None or collapse_b_const is None):
        collapse_a_const = member_risk_rows[0].get("a")
        collapse_b_const = member_risk_rows[0].get("b")
    elif node_risk_rows and (collapse_a_const is None or collapse_b_const is None):
        collapse_a_const = node_risk_rows[0].get("a")
        collapse_b_const = node_risk_rows[0].get("b")
    if member_risk_rows:
        collapse_vr_const = member_risk_rows[0].get("vr")
    elif node_risk_rows:
        collapse_vr_const = node_risk_rows[0].get("vr")

    collapse_member_rows_override, collapse_joint_rows_override = build_collapse_overrides_from_sheet(
        ws_collapse=ws_collapse,
        ws_members=ws_members,
        node_type_by_joint=node_type_by_joint,
        member_level_by_pair=member_level_by_pair,
        joint_level_by_id=joint_level_by_id,
        collapse_a_const=collapse_a_const,
        collapse_b_const=collapse_b_const,
        collapse_vr_const=collapse_vr_const,
    )

    return build_context(
        node_risk_rows,
        node_strategy_rows,
        member_risk_rows,
        member_strategy_rows,
        metadata,
        row_limits=row_limits,
        collapse_member_rows_override=collapse_member_rows_override,
        collapse_joint_rows_override=collapse_joint_rows_override,
        apply_vba_member_delete_rules=apply_vba_member_delete_rules,
        apply_vba_joint_delete_rules_current=apply_vba_joint_delete_rules_current,
        apply_vba_joint_delete_rules_future=apply_vba_joint_delete_rules_future,
    )


def load_context_from_python_workbook(
    wb,
    metadata: dict[str, Any],
    row_limits: dict[str, int] | None = None,
    strict_vba_algorithms: bool = True,
) -> dict[str, Any]:
    ws_node_risk = wb["\u8282\u70b9\u5931\u6548\u98ce\u9669\u7b49\u7ea7(Python)"]
    ws_forecast = wb["\u8282\u70b9\u5931\u6548\u98ce\u9669\u7b49\u7ea7\u9884\u6d4b(\u4ec5\u75b2\u52b3Python)"]
    ws_node_plan = wb["\u8282\u70b9\u68c0\u9a8c\u7b56\u7565(Python)"]
    ws_fatigue = wb["\u75b2\u52b3\u5206\u6790\u7ed3\u679c"] if "\u75b2\u52b3\u5206\u6790\u7ed3\u679c" in wb.sheetnames else None
    ws_members = wb["Members"] if "Members" in wb.sheetnames else None
    ws_collapse = wb["\u5012\u584c\u5206\u6790\u7ed3\u679c"] if "\u5012\u584c\u5206\u6790\u7ed3\u679c" in wb.sheetnames else None

    brace_by_joint: dict[str, tuple[str, float]] = {}
    if ws_fatigue is not None:
        for row in iter_data_rows(ws_fatigue, min_row=2, max_col=13):
            joint = to_text(row[0])
            brace = to_text(row[1])
            dmax = to_float_or_none(row[12]) or 0.0
            if not joint:
                continue
            old = brace_by_joint.get(joint)
            if old is None or dmax > old[1]:
                brace_by_joint[joint] = (brace, dmax)

    node_risk_rows: list[dict[str, Any]] = []
    node_base_by_joint: dict[str, dict[str, Any]] = {}
    for row in iter_data_rows(ws_node_risk, min_row=2, max_col=15):
        joint_id = to_text(row[0])
        if joint_id == "":
            continue
        one = {
            "joint_id": joint_id,
            "brace": brace_by_joint.get(joint_id, ("", 0.0))[0],
            "joint_type": to_text(row[1]),
            "consequence_level": "3",
            "a": to_text(row[2]),
            "b": to_text(row[3]),
            "rm": to_text(row[4]),
            "vr": to_text(row[5]),
            "pf": to_text(row[6]),
            "collapse_prob_level": to_text(row[7]),
            "damage": to_text(row[8]),
            "beta": to_text(row[10]),
            "pf_fatigue": to_text(row[11]),
            "fatigue_prob_level": to_text(row[12]),
            "combined_prob_level": to_text(row[13]),
            "node_risk_level": to_text(row[14]),
        }
        node_risk_rows.append(one)
        node_base_by_joint[joint_id] = one

    collapse_a_const = None
    collapse_b_const = None
    collapse_vr_const = None
    if node_risk_rows:
        collapse_a_const = node_risk_rows[0].get("a")
        collapse_b_const = node_risk_rows[0].get("b")
        collapse_vr_const = node_risk_rows[0].get("vr")

    node_plan_map: dict[tuple[str, str], dict[str, Any]] = {}
    for row in iter_data_rows(ws_node_plan, min_row=2, max_col=5):
        joint_id = to_text(row[0])
        time_node = normalize_plan_time(row[1])
        if joint_id == "" or time_node == "":
            continue
        node_plan_map[(joint_id, time_node)] = {
            "inspect_level": normalize_inspect_level(row[3]),
            "risk_grade": to_text(row[2]),
        }

    node_strategy_rows: list[dict[str, Any]] = []
    for base in node_risk_rows:
        joint_id = base["joint_id"]
        plan_cur = node_plan_map.get((joint_id, TIME_CURRENT), {})
        node_strategy_rows.append(
            {
                "joint_id": joint_id,
                "brace": base["brace"],
                "joint_type": base["joint_type"],
                "consequence_level": base["consequence_level"],
                "collapse_prob_level": base["collapse_prob_level"],
                "damage": to_scientific_text(base["damage"]),
                "beta": base["beta"],
                "pf": to_scientific_text(base["pf_fatigue"]),
                "fatigue_prob_level": base["fatigue_prob_level"],
                "combined_prob_level": base["combined_prob_level"],
                "node_risk_level": base["node_risk_level"],
                "inspect_level": plan_cur.get("inspect_level", "III"),
                "time_node": TIME_CURRENT,
            }
        )

    forecast_headers = [to_text(ws_forecast.cell(row=1, column=c).value) for c in range(1, min(ws_forecast.max_column, 120) + 1)]
    is_long_forecast = ("Year" in forecast_headers and "D_future" in forecast_headers) or (
        len(forecast_headers) >= 7 and forecast_headers[:7] == ["JoitID", "Year", "D_future", "beta", "Pf", "PossLevel", "RiskGrade"]
    )

    if is_long_forecast:
        for row in iter_data_rows(ws_forecast, min_row=2, max_col=7):
            joint_id = to_text(row[0])
            time_node = forecast_year_to_time_node(to_int_or_none(row[1]))
            if joint_id == "" or time_node == "" or time_node == TIME_CURRENT:
                continue

            base = node_base_by_joint.get(joint_id, {})
            collapse_lv = to_int_or_none(base.get("collapse_prob_level"))
            fatigue_lv = to_int_or_none(row[5])
            if collapse_lv is None:
                combined_lv = fatigue_lv
            elif fatigue_lv is None:
                combined_lv = collapse_lv
            else:
                combined_lv = min(collapse_lv, fatigue_lv)

            plan = node_plan_map.get((joint_id, time_node), {})
            node_strategy_rows.append(
                {
                    "joint_id": joint_id,
                    "brace": to_text(base.get("brace")),
                    "joint_type": to_text(base.get("joint_type")),
                    "consequence_level": to_text(base.get("consequence_level")) or "3",
                    "collapse_prob_level": to_text(base.get("collapse_prob_level")),
                    "damage": to_scientific_text(row[2]),
                    "beta": to_text(row[3]),
                    "pf": to_scientific_text(row[4]),
                    "fatigue_prob_level": to_text(row[5]),
                    "combined_prob_level": "" if combined_lv is None else str(combined_lv),
                    "node_risk_level": to_text(row[6]) or to_text(base.get("node_risk_level")),
                    "inspect_level": plan.get("inspect_level", "III"),
                    "time_node": time_node,
                }
            )
    else:
        # VBA-wide forecast sheet:
        # [JoitID, Brace, JointType, ConsequenceLevel, CollapsePossLevel,
        #  N_*, N+5_*, N+10_*, N+15_*, N+20_*, N+25_*]
        horizon_blocks = [
            ("N", TIME_CURRENT, 5),
            ("N+5", "第5年", 16),
            ("N+10", "第10年", 27),
            ("N+15", "第15年", 38),
            ("N+20", "第20年", 49),
            ("N+25", "第25年", 60),
        ]
        for row in iter_data_rows(ws_forecast, min_row=2, max_col=71):
            joint_id = to_text(row[0])
            if joint_id == "":
                continue
            base = node_base_by_joint.get(joint_id, {})
            collapse_lv = to_int_or_none(base.get("collapse_prob_level"))

            for _, time_node, start in horizon_blocks:
                if time_node == TIME_CURRENT:
                    continue
                if len(row) <= start + 10:
                    continue

                damage = to_scientific_text(row[start + 0])
                beta = to_text(row[start + 6])
                pf = to_scientific_text(row[start + 7])
                fatigue_text = to_text(row[start + 8])
                combined_text = to_text(row[start + 9])
                risk_text = to_text(row[start + 10])

                # Skip empty block rows.
                if damage == "" and beta == "" and pf == "" and fatigue_text == "" and combined_text == "" and risk_text == "":
                    continue

                fatigue_lv = to_int_or_none(fatigue_text)
                combined_lv = to_int_or_none(combined_text)
                if combined_lv is None:
                    if collapse_lv is None:
                        combined_lv = fatigue_lv
                    elif fatigue_lv is None:
                        combined_lv = collapse_lv
                    else:
                        combined_lv = min(collapse_lv, fatigue_lv)

                plan = node_plan_map.get((joint_id, time_node), {})
                node_strategy_rows.append(
                    {
                        "joint_id": joint_id,
                        "brace": to_text(base.get("brace")) or to_text(row[1]),
                        "joint_type": to_text(base.get("joint_type")) or to_text(row[2]),
                        "consequence_level": to_text(base.get("consequence_level")) or to_text(row[3]) or "3",
                        "collapse_prob_level": to_text(base.get("collapse_prob_level")) or to_text(row[4]),
                        "damage": damage,
                        "beta": beta,
                        "pf": pf,
                        "fatigue_prob_level": fatigue_text,
                        "combined_prob_level": "" if combined_lv is None else str(combined_lv),
                        "node_risk_level": risk_text or to_text(base.get("node_risk_level")),
                        "inspect_level": plan.get("inspect_level", "III"),
                        "time_node": time_node,
                    }
                )

    member_type_by_pair: dict[str, str] = {}
    collapse_member_factor: dict[str, float | None] = {}
    collapse_joint_factor: dict[str, float | None] = {}
    if ws_collapse is not None:
        for row in iter_data_rows(ws_collapse, min_row=2, max_col=6, key_col_index=1):
            fail_type = to_text(row[1])
            location = to_text(row[2])
            factor = to_float_or_none(row[3])
            if fail_type == "\u6784\u4ef6\u5931\u6548" and location:
                old = collapse_member_factor.get(location)
                if old is None or (factor is not None and factor < old):
                    collapse_member_factor[location] = factor
            elif fail_type == "\u8282\u70b9\u5931\u6548" and location:
                old = collapse_joint_factor.get(location)
                if old is None or (factor is not None and factor < old):
                    collapse_joint_factor[location] = factor

    member_risk_rows: list[dict[str, Any]] = []
    seen_member_pairs: set[tuple[str, str]] = set()
    if ws_members is not None:
        for row in iter_data_rows(ws_members, min_row=2, max_col=7):
            joint_a = to_text(row[0])
            joint_b = to_text(row[1])
            if joint_a == "" or joint_b == "":
                continue
            pair = (joint_a, joint_b)
            if pair in seen_member_pairs:
                continue
            seen_member_pairs.add(pair)

            pair_key = f"{joint_a}-{joint_b}"
            rev_key = f"{joint_b}-{joint_a}"
            mtype = to_text(row[4]) or "Other"
            if pair_key not in member_type_by_pair:
                member_type_by_pair[pair_key] = mtype
            if rev_key not in member_type_by_pair:
                member_type_by_pair[rev_key] = mtype

            base_a = node_base_by_joint.get(joint_a, {})
            base_b = node_base_by_joint.get(joint_b, {})

            collapse_a = to_int_or_none(base_a.get("collapse_prob_level"))
            collapse_b = to_int_or_none(base_b.get("collapse_prob_level"))
            if collapse_a is None:
                collapse_lv = collapse_b
            elif collapse_b is None:
                collapse_lv = collapse_a
            else:
                collapse_lv = min(collapse_a, collapse_b)

            pf_a = to_float_or_none(base_a.get("pf"))
            pf_b = to_float_or_none(base_b.get("pf"))
            if pf_a is None:
                pf_member = pf_b
            elif pf_b is None:
                pf_member = pf_a
            else:
                pf_member = min(pf_a, pf_b)

            member_risk = choose_higher_risk_grade(base_a.get("node_risk_level"), base_b.get("node_risk_level"))
            if member_risk == "":
                member_risk = collapse_level_to_risk_grade(collapse_lv) or "\u56db"

            factor = collapse_member_factor.get(pair_key)
            if factor is None:
                factor = collapse_member_factor.get(rev_key)

            if factor is not None:
                pf_calc = collapse_pf_from_factor(collapse_a_const, collapse_b_const, factor, collapse_vr_const)
                rm_val = f"{factor:g}"
                pf_val = "" if pf_calc is None else f"{pf_calc:g}"
                collapse_level = "1"
            else:
                rm_val = to_text(base_a.get("rm")) or to_text(base_b.get("rm"))
                pf_val = "" if pf_member is None else f"{pf_member:g}"
                collapse_level = "" if collapse_lv is None else str(collapse_lv)

            member_risk_rows.append(
                {
                    "joint_a": joint_a,
                    "joint_b": joint_b,
                    "member_type": mtype,
                    "consequence_level": "3",
                    "a": to_text(base_a.get("a")) or to_text(base_b.get("a")),
                    "b": to_text(base_a.get("b")) or to_text(base_b.get("b")),
                    "rm": rm_val,
                    "vr": to_text(base_a.get("vr")) or to_text(base_b.get("vr")),
                    "pf": pf_val,
                    "collapse_prob_level": collapse_level,
                    "member_risk_level": member_risk,
                }
            )

    member_level_policy = {"\u4e00": "II", "\u4e8c": "II", "\u4e09": "III", "\u56db": "IV", "\u4e94": "IV"}
    member_strategy_rows: list[dict[str, Any]] = []
    for r in member_risk_rows:
        member_strategy_rows.append(
            {
                "joint_a": r["joint_a"],
                "joint_b": r["joint_b"],
                "member_type": r["member_type"],
                "consequence_level": r["consequence_level"],
                "member_risk_level": r["member_risk_level"],
                "inspect_level": member_level_policy.get(to_text(r["member_risk_level"]), "III"),
                "time_node": TIME_CURRENT,
            }
        )

    collapse_member_rows_override: list[dict[str, Any]] = []
    for location, factor in sorted(
        collapse_member_factor.items(),
        key=lambda kv: (
            9999 if kv[1] is None else kv[1],
            kv[0],
        ),
    ):
        joint_a = ""
        joint_b = ""
        if "-" in location:
            joint_a, joint_b = location.split("-", 1)
        mtype = member_type_by_pair.get(location) or member_type_by_pair.get(f"{joint_b}-{joint_a}") or "Other"
        pf_val = collapse_pf_from_factor(collapse_a_const, collapse_b_const, factor, collapse_vr_const)
        collapse_member_rows_override.append(
            {
                "joint_a": joint_a,
                "joint_b": joint_b,
                "member_type": mtype,
                "a": to_text(collapse_a_const),
                "b": to_text(collapse_b_const),
                "rm": "" if factor is None else f"{factor:g}",
                "vr": to_text(collapse_vr_const),
                "pf": "" if pf_val is None else f"{pf_val:g}",
                "collapse_prob_level": "1",
            }
        )

    collapse_joint_rows_override: list[dict[str, Any]] = []
    for joint_id, factor in sorted(
        collapse_joint_factor.items(),
        key=lambda kv: (
            9999 if kv[1] is None else kv[1],
            kv[0],
        ),
    ):
        base = node_base_by_joint.get(joint_id, {})
        pf_val = collapse_pf_from_factor(collapse_a_const, collapse_b_const, factor, collapse_vr_const)
        collapse_joint_rows_override.append(
            {
                "joint_id": joint_id,
                "joint_type": to_text(base.get("joint_type")),
                "a": to_text(collapse_a_const),
                "b": to_text(collapse_b_const),
                "rm": "" if factor is None else f"{factor:g}",
                "vr": to_text(collapse_vr_const),
                "pf": "" if pf_val is None else f"{pf_val:g}",
                "collapse_prob_level": to_text(base.get("collapse_prob_level")),
            }
        )

    collapse_member_override = None
    collapse_joint_override = None
    # Do not inject fallback/override rows here.
    # Report tables must be driven by real computed risk sheets.

    return build_context(
        node_risk_rows,
        node_strategy_rows,
        member_risk_rows,
        member_strategy_rows,
        metadata,
        row_limits=row_limits,
        collapse_member_rows_override=collapse_member_override,
        collapse_joint_rows_override=collapse_joint_override,
        apply_vba_member_delete_rules=strict_vba_algorithms,
        apply_vba_joint_delete_rules=strict_vba_algorithms,
    )


def load_context_from_workbook(
    workbook_path: Path,
    metadata: dict[str, Any],
    row_limits: dict[str, int] | None = None,
    strict_vba_algorithms: bool = True,
    apply_vba_member_delete_rules: bool = False,
    apply_vba_joint_delete_rules_current: bool = False,
    apply_vba_joint_delete_rules_future: bool = False,
) -> dict[str, Any]:
    wb = load_workbook(workbook_path, data_only=True, read_only=True)

    # Strict mode: always read template sheets directly (legacy-style).
    # This keeps report output aligned with VBA sheet results and avoids
    # secondary reconstruction differences.
    if strict_vba_algorithms:
        return load_context_from_legacy_workbook(
            wb,
            metadata,
            row_limits=row_limits,
            apply_vba_member_delete_rules=apply_vba_member_delete_rules,
            apply_vba_joint_delete_rules_current=apply_vba_joint_delete_rules_current,
            apply_vba_joint_delete_rules_future=apply_vba_joint_delete_rules_future,
        )

    python_sheet_names = {
        "\u8282\u70b9\u5931\u6548\u98ce\u9669\u7b49\u7ea7(Python)",
        "\u8282\u70b9\u5931\u6548\u98ce\u9669\u7b49\u7ea7\u9884\u6d4b(\u4ec5\u75b2\u52b3Python)",
        "\u8282\u70b9\u68c0\u9a8c\u7b56\u7565(Python)",
    }
    if python_sheet_names.issubset(set(wb.sheetnames)):
        return load_context_from_python_workbook(
            wb,
            metadata,
            row_limits=row_limits,
            strict_vba_algorithms=strict_vba_algorithms,
        )
    return load_context_from_legacy_workbook(
        wb,
        metadata,
        row_limits=row_limits,
        apply_vba_member_delete_rules=apply_vba_member_delete_rules,
        apply_vba_joint_delete_rules_current=apply_vba_joint_delete_rules_current,
        apply_vba_joint_delete_rules_future=apply_vba_joint_delete_rules_future,
    )


def list_appendix_files(path: Path) -> list[Path]:
    return sorted((p for p in path.rglob("*") if p.is_file()), key=lambda p: (str(p.parent), p.name.lower()))


def _text_entry(text: str) -> dict[str, Any]:
    return {"kind": "text", "text": text}


def _link_entry(text: str, target: Path) -> dict[str, Any]:
    return {"kind": "link", "text": text, "target": target.resolve().as_uri()}


def build_appendix_sections(
    appendix_a_file: str = "",
    appendix_b_file: str = "",
    appendix_c_dirs: list[str] | None = None,
) -> list[dict[str, Any]]:
    appendix_c_dirs = appendix_c_dirs or []

    def build_file_section(heading: str, path_text: str) -> dict[str, Any]:
        entries: list[dict[str, Any]] = []
        if not path_text.strip():
            entries.append(_text_entry("未配置文件路径。"))
        else:
            path = Path(path_text).resolve()
            if path.exists() and path.is_file():
                entries.append(_text_entry(f"来源文件：{path}"))
                entries.append(_link_entry(f"打开文件：{path.name}", path))
            else:
                entries.append(_text_entry(f"文件不存在：{path}"))
        return {"heading": heading, "entries": entries}

    sections = [
        build_file_section(APPENDIX_A_HEADING, appendix_a_file),
        build_file_section(APPENDIX_B_HEADING, appendix_b_file),
    ]

    summary_entries: list[dict[str, Any]] = []
    if appendix_c_dirs:
        summary_entries.append(_text_entry("以下附件C子项按配置目录自动读取该目录下的全部文件。"))
        for idx, path_text in enumerate(appendix_c_dirs[: len(APPENDIX_C_SUBHEADINGS)], start=1):
            path = Path(path_text).resolve()
            if not path_text.strip():
                summary_entries.append(_text_entry(f"C.{idx}：未配置目录路径。"))
                continue
            if not path.exists() or not path.is_dir():
                summary_entries.append(_text_entry(f"C.{idx}：目录不存在：{path}"))
                continue
            files = list_appendix_files(path)
            summary_entries.append(_text_entry(f"C.{idx}：{path}（{len(files)} 个文件）"))
    else:
        summary_entries.append(_text_entry("未配置附件C目录。"))
    sections.append({"heading": APPENDIX_C_MAIN_HEADING, "entries": summary_entries})

    for heading, path_text in zip(APPENDIX_C_SUBHEADINGS, appendix_c_dirs):
        entries: list[dict[str, Any]] = []
        if not path_text.strip():
            entries.append(_text_entry("未配置目录路径。"))
            sections.append({"heading": heading, "entries": entries})
            continue
        path = Path(path_text).resolve()
        if not path.exists() or not path.is_dir():
            entries.append(_text_entry(f"目录不存在：{path}"))
            sections.append({"heading": heading, "entries": entries})
            continue
        files = list_appendix_files(path)
        entries.append(_text_entry(f"来源目录：{path}"))
        entries.append(_text_entry(f"文件数量：{len(files)}"))
        if files:
            for idx, file_path in enumerate(files, start=1):
                entries.append(_link_entry(f"{idx}. {file_path.name}", file_path))
        else:
            entries.append(_text_entry("该目录下未读取到文件。"))
        sections.append({"heading": heading, "entries": entries})

    for heading in APPENDIX_C_SUBHEADINGS[len(appendix_c_dirs) :]:
        sections.append({"heading": heading, "entries": [_text_entry("未配置目录路径。")]})

    return sections


def build_appendix_pdf_plan(
    appendix_a_file: str = "",
    appendix_b_file: str = "",
    appendix_c_dirs: list[str] | None = None,
) -> list[dict[str, Any]]:
    appendix_c_dirs = appendix_c_dirs or []

    sections: list[dict[str, Any]] = []
    for heading, path_text in [
        (APPENDIX_A_HEADING, appendix_a_file),
        (APPENDIX_B_HEADING, appendix_b_file),
    ]:
        files: list[Path] = []
        if path_text.strip():
            path = Path(path_text).resolve()
            if path.exists() and path.is_file():
                files.append(path)
        sections.append({"heading": heading, "files": files})

    sections.append({"heading": APPENDIX_C_MAIN_HEADING, "files": []})
    for heading, path_text in zip(APPENDIX_C_SUBHEADINGS, appendix_c_dirs):
        files: list[Path] = []
        if path_text.strip():
            path = Path(path_text).resolve()
            if path.exists() and path.is_dir():
                files = list_appendix_files(path)
        sections.append({"heading": heading, "files": files})

    for heading in APPENDIX_C_SUBHEADINGS[len(appendix_c_dirs) :]:
        sections.append({"heading": heading, "files": []})

    return sections


def export_workbook_sheets_to_pdf(
    workbook_path: Path,
    sheet_exports: list[tuple[str, str, Path]],
) -> list[Path]:
    target = Path(workbook_path).resolve()
    if not target.exists():
        raise FileNotFoundError(f"Workbook not found: {target}")
    if not sheet_exports:
        return []

    excel = None
    workbook = None
    exported: list[Path] = []
    pythoncom.CoInitialize()
    try:
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = 0
        workbook = excel.Workbooks.Open(
            str(target),
            False,
            True,
        )
        for _, sheet_name, output_pdf in sheet_exports:
            try:
                worksheet = workbook.Worksheets(sheet_name)
            except Exception as exc:
                raise KeyError(f"Appendix worksheet not found: {sheet_name}") from exc

            output_pdf = Path(output_pdf).resolve()
            output_pdf.parent.mkdir(parents=True, exist_ok=True)
            if output_pdf.exists():
                output_pdf.unlink(missing_ok=True)

            page_setup = worksheet.PageSetup
            page_setup.Orientation = 2  # xlLandscape
            page_setup.Zoom = False
            page_setup.FitToPagesWide = 1
            page_setup.FitToPagesTall = False
            worksheet.ExportAsFixedFormat(0, str(output_pdf))  # xlTypePDF
            exported.append(output_pdf)
        return exported
    finally:
        if workbook is not None:
            try:
                workbook.Close(SaveChanges=False)
            except Exception:
                pass
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def build_generated_appendix_plan(
    workbook_path: Path,
    *,
    facility_code: str = "",
    run_id: int | None = None,
    scratch_dir: Path | None = None,
) -> tuple[list[dict[str, Any]], list[Path]]:
    workbook = Path(workbook_path).resolve()
    scratch_root = Path(scratch_dir or workbook.parent).resolve()
    scratch_root.mkdir(parents=True, exist_ok=True)

    plan: list[dict[str, Any]] = []
    temp_files: list[Path] = []

    sheet_exports = [
        (heading, sheet_name, (scratch_root / filename).resolve())
        for heading, sheet_name, filename in GENERATED_APPENDIX_SHEETS
    ]
    temp_files.extend(export_workbook_sheets_to_pdf(workbook, sheet_exports))
    for heading, _, output_pdf in sheet_exports:
        plan.append({"heading": heading, "files": [output_pdf] if output_pdf.exists() else []})

    plan.append({"heading": APPENDIX_C_MAIN_HEADING, "files": []})

    if facility_code.strip():
        from services.special_strategy_image_service import build_strategy_image_path
        try:
            from services.special_strategy_state_db import list_strategy_risk_images
        except Exception:
            list_strategy_risk_images = None

        def normalize_year_key(value: object) -> str:
            text = str(value or "").strip()
            match = re.search(r"(\d+)", text)
            if match:
                return match.group(1)
            return text.replace(" ", "").replace("+", "")

        def sort_key_from_name(name: str) -> tuple[int, str]:
            stem = str(name or "").replace("_", " ").strip()
            compact = stem.replace(" ", "")
            for idx, row_name in enumerate(APPENDIX_C_ROW_NAMES):
                row_compact = row_name.replace(" ", "")
                if row_compact in compact:
                    return idx, stem.lower()
            return len(APPENDIX_C_ROW_NAMES), stem.lower()

        def sort_key(image_path: Path) -> tuple[int, str]:
            return sort_key_from_name(image_path.stem)

        def caption_from_path(image_path: Path) -> str:
            stem = image_path.stem.replace("_", " ").strip()
            compact = stem.replace(" ", "")
            for row_name in APPENDIX_C_ROW_NAMES:
                if row_name.replace(" ", "") in compact:
                    return row_name
            return stem

        def list_year_images_from_records(target_run_id: int | None, year_label: str) -> list[dict[str, Any]]:
            if list_strategy_risk_images is None:
                return []
            try:
                kwargs: dict[str, Any] = {
                    "page_code": "upgrade_special_inspection_result",
                    "limit": 5000,
                }
                if target_run_id is not None:
                    kwargs["run_id"] = target_run_id
                records = list_strategy_risk_images(facility_code, **kwargs)
            except Exception:
                return []

            expected_year = normalize_year_key(year_label)
            collected: list[dict[str, Any]] = []
            seen_paths: set[str] = set()
            for record in records:
                if normalize_year_key(record.get("year_label")) != expected_year:
                    continue
                image_path = Path(str(record.get("image_path") or "")).expanduser()
                if not image_path.exists() or not image_path.is_file():
                    continue
                key = str(image_path).lower()
                if key in seen_paths:
                    continue
                seen_paths.add(key)
                caption = str(record.get("row_name") or "").strip() or caption_from_path(image_path)
                collected.append(
                    {
                        "path": image_path,
                        "caption": caption,
                        "_sort_name": caption,
                    }
                )

            collected.sort(key=lambda item: sort_key_from_name(str(item.get("_sort_name") or "")))
            for item in collected:
                item.pop("_sort_name", None)
            return collected

        def list_year_images(target_run_id: int | None, year_label: str) -> list[Path]:
            year_dir = build_strategy_image_path(
                facility_code=facility_code,
                run_id=target_run_id,
                page_code="upgrade_special_inspection_result",
                image_type="elevation_risk",
                year_label=year_label,
                row_name="__probe__",
                create_dirs=False,
            ).parent
            if not year_dir.exists() or not year_dir.is_dir():
                return []
            return sorted(
                [p for p in year_dir.iterdir() if p.is_file() and p.suffix.lower() in APPENDIX_IMAGE_SUFFIXES],
                key=sort_key,
            )

        for year_label, heading in APPENDIX_C_YEAR_SECTION_MAP:
            section_files = list_year_images_from_records(run_id, year_label)
            if not section_files and run_id is not None:
                section_files = list_year_images_from_records(None, year_label)
            if not section_files:
                image_paths = list_year_images(run_id, year_label)
                if not image_paths and run_id is not None:
                    image_paths = list_year_images(None, year_label)
                section_files = [
                    {"path": image_path, "caption": caption_from_path(image_path)}
                    for image_path in image_paths
                ]
            plan.append({"heading": heading, "files": section_files})
    else:
        for _, heading in APPENDIX_C_YEAR_SECTION_MAP:
            plan.append({"heading": heading, "files": []})

    return plan, temp_files


def get_table_list(document_root: ET.Element) -> list[ET.Element]:
    body = document_root.find("w:body", NS)
    if body is None:
        return []
    return [child for child in body if child.tag.endswith("tbl")]


def ensure_update_fields_on_open(settings_root: ET.Element) -> None:
    node = settings_root.find("w:updateFields", NS)
    if node is None:
        node = ET.SubElement(settings_root, f"{{{NS['w']}}}updateFields")
    node.set(f"{{{NS['w']}}}val", "true")


def _paragraph_style_value(paragraph: ET.Element) -> str:
    style = paragraph.find("./w:pPr/w:pStyle", NS)
    if style is None:
        return ""
    return style.get(f"{{{NS['w']}}}val", "")


def _paragraph_has_field(paragraph: ET.Element, keyword: str) -> bool:
    instr_nodes = paragraph.findall(".//w:instrText", NS)
    if not instr_nodes:
        return False
    return any(keyword in "".join(node.itertext()) for node in instr_nodes)


def _is_toc_related_paragraph(paragraph: ET.Element) -> bool:
    style_value = _paragraph_style_value(paragraph)
    if style_value.startswith("TOC"):
        return True
    if _paragraph_has_field(paragraph, "TOC"):
        return True
    if _paragraph_has_field(paragraph, "PAGEREF _Toc"):
        return True
    return False


def create_toc_field_paragraph() -> ET.Element:
    paragraph = ET.Element(f"{{{NS['w']}}}p")
    paragraph_pr = ET.SubElement(paragraph, f"{{{NS['w']}}}pPr")
    ET.SubElement(paragraph_pr, f"{{{NS['w']}}}spacing", {
        f"{{{NS['w']}}}after": "0",
        f"{{{NS['w']}}}line": "300",
        f"{{{NS['w']}}}lineRule": "auto",
    })

    run_begin = ET.SubElement(paragraph, f"{{{NS['w']}}}r")
    ET.SubElement(
        run_begin,
        f"{{{NS['w']}}}fldChar",
        {
            f"{{{NS['w']}}}fldCharType": "begin",
            f"{{{NS['w']}}}dirty": "true",
        },
    )

    run_instr = ET.SubElement(paragraph, f"{{{NS['w']}}}r")
    instr = ET.SubElement(run_instr, f"{{{NS['w']}}}instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    run_separate = ET.SubElement(paragraph, f"{{{NS['w']}}}r")
    ET.SubElement(run_separate, f"{{{NS['w']}}}fldChar", {f"{{{NS['w']}}}fldCharType": "separate"})

    placeholder_run = ET.SubElement(paragraph, f"{{{NS['w']}}}r")
    placeholder_text = ET.SubElement(placeholder_run, f"{{{NS['w']}}}t")
    placeholder_text.text = "目录将在打开文档后自动更新"

    run_end = ET.SubElement(paragraph, f"{{{NS['w']}}}r")
    ET.SubElement(run_end, f"{{{NS['w']}}}fldChar", {f"{{{NS['w']}}}fldCharType": "end"})
    return paragraph


def rebuild_toc_field(document_root: ET.Element) -> None:
    body = document_root.find("w:body", NS)
    if body is None:
        return

    children = list(body)
    heading_idx = -1
    for idx, child in enumerate(children):
        if not child.tag.endswith("p"):
            continue
        if normalize_lookup_text(paragraph_text(child)) == "目录":
            heading_idx = idx
            break
    if heading_idx < 0:
        return

    remove_targets: list[ET.Element] = []
    for child in children[heading_idx + 1 :]:
        if not child.tag.endswith("p"):
            break
        if not _is_toc_related_paragraph(child):
            break
        remove_targets.append(child)

    for paragraph in remove_targets:
        body.remove(paragraph)

    toc_field = create_toc_field_paragraph()
    insert_pos = heading_idx + 1
    body.insert(insert_pos, toc_field)


def sanitize_markup_compatibility(document_root: ET.Element) -> None:
    mc_uri = WORD_NAMESPACE_PREFIXES["mc"]
    ignorable_key = f"{{{mc_uri}}}Ignorable"
    ignorable_value = str(document_root.attrib.get(ignorable_key, "") or "").strip()
    if not ignorable_value:
        return

    used_uris: set[str] = set()
    for element in document_root.iter():
        if element.tag.startswith("{"):
            used_uris.add(element.tag.split("}", 1)[0][1:])
        for attr_name in element.attrib:
            if attr_name.startswith("{"):
                used_uris.add(attr_name.split("}", 1)[0][1:])

    filtered_prefixes = [
        prefix
        for prefix in ignorable_value.split()
        if WORD_NAMESPACE_PREFIXES.get(prefix) in used_uris
    ]
    if filtered_prefixes:
        document_root.set(ignorable_key, " ".join(filtered_prefixes))
    else:
        document_root.attrib.pop(ignorable_key, None)


def mark_toc_fields_dirty(document_root: ET.Element) -> None:
    for paragraph in document_root.findall(".//w:p", NS):
        instr_texts = paragraph.findall(".//w:instrText", NS)
        if not instr_texts:
            continue
        if not any("TOC" in "".join(node.itertext()) for node in instr_texts):
            continue
        for field in paragraph.findall(".//w:fldChar", NS):
            if field.get(f"{{{NS['w']}}}fldCharType") == "begin":
                field.set(f"{{{NS['w']}}}dirty", "true")


def refresh_word_document_fields(
    docx_path: Path,
    timeout_seconds: int = 90,
    pdf_output_path: Path | None = None,
) -> bool:
    target = Path(docx_path).resolve()
    if not target.exists():
        return False
    pdf_target = Path(pdf_output_path).resolve() if pdf_output_path else None
    if pdf_target is not None:
        pdf_target.parent.mkdir(parents=True, exist_ok=True)
    word = None
    document = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            FileName=str(target),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        document.Repaginate()
        toc_count = int(getattr(document.TablesOfContents, "Count", 0) or 0)
        for idx in range(1, toc_count + 1):
            document.TablesOfContents(idx).Update()
        document.Fields.Update()
        document.Save()
        if pdf_target is not None:
            document.ExportAsFixedFormat(
                OutputFileName=str(pdf_target),
                ExportFormat=17,  # wdExportFormatPDF
            )
        return True
    except Exception:
        return False
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def normalize_lookup_text(value: Any) -> str:
    return re.sub(r"\s+", "", to_text(value))


def _table_header_lookup_text(table: ET.Element) -> str:
    first_row = table.find("./w:tr", NS)
    if first_row is None:
        return ""
    return normalize_lookup_text("".join(t.text or "" for t in first_row.findall(".//w:t", NS)))


def find_table_by_locator(document_root: ET.Element, locator: TableLocator) -> ET.Element | None:
    body = document_root.find("w:body", NS)
    if body is not None:
        last_paragraph = ""
        title_matches: list[tuple[ET.Element, str]] = []
        for child in body:
            if child.tag.endswith("p"):
                text = normalize_lookup_text(paragraph_text(child))
                if text:
                    last_paragraph = text
                continue
            if not child.tag.endswith("tbl"):
                continue
            title_match = (
                not locator.title_keywords
                or all(normalize_lookup_text(key) in last_paragraph for key in locator.title_keywords)
            )
            if not title_match:
                continue
            title_matches.append((child, _table_header_lookup_text(child)))

        if title_matches:
            if not locator.header_keywords:
                return title_matches[0][0]

            normalized_headers = tuple(normalize_lookup_text(key) for key in locator.header_keywords)
            for table, header_text in title_matches:
                if all(key in header_text for key in normalized_headers):
                    return table
            return title_matches[0][0]

    tables = get_table_list(document_root)
    if locator.fallback_index is not None and 0 <= locator.fallback_index < len(tables):
        return tables[locator.fallback_index]
    return None


def get_text_nodes(el: ET.Element) -> list[ET.Element]:
    return el.findall(".//w:t", NS)


def set_cell_text(tc: ET.Element, value: str) -> None:
    text_nodes = get_text_nodes(tc)
    if not text_nodes:
        p = ET.SubElement(tc, f"{{{NS['w']}}}p")
        r = ET.SubElement(p, f"{{{NS['w']}}}r")
        t = ET.SubElement(r, f"{{{NS['w']}}}t")
        t.text = value
        return
    text_nodes[0].text = value
    for t in text_nodes[1:]:
        t.text = ""


def set_paragraph_text(p: ET.Element, value: str) -> None:
    text_nodes = get_text_nodes(p)
    if not text_nodes:
        r = ET.SubElement(p, f"{{{NS['w']}}}r")
        t = ET.SubElement(r, f"{{{NS['w']}}}t")
        t.text = value
        return
    text_nodes[0].text = value
    for t in text_nodes[1:]:
        t.text = ""


def paragraph_text(p: ET.Element) -> str:
    return "".join((t.text or "") for t in p.findall(".//w:t", NS)).strip()


def create_simple_paragraph(value: str) -> ET.Element:
    p = ET.Element(f"{{{NS['w']}}}p")
    r = ET.SubElement(p, f"{{{NS['w']}}}r")
    t = ET.SubElement(r, f"{{{NS['w']}}}t")
    if value.startswith(" ") or value.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = value
    return p


def _next_relationship_id(rels_root: ET.Element) -> str:
    max_idx = 0
    for rel in rels_root.findall(f"{{{REL_NS}}}Relationship"):
        rel_id = rel.get("Id", "")
        if rel_id.startswith("rId"):
            suffix = rel_id[3:]
            if suffix.isdigit():
                max_idx = max(max_idx, int(suffix))
    return f"rId{max_idx + 1}"


def create_hyperlink_paragraph(text: str, target: str, rels_root: ET.Element) -> ET.Element:
    rel_id = _next_relationship_id(rels_root)
    ET.SubElement(
        rels_root,
        f"{{{REL_NS}}}Relationship",
        {
            "Id": rel_id,
            "Type": HYPERLINK_REL_TYPE,
            "Target": target,
            "TargetMode": "External",
        },
    )

    p = ET.Element(f"{{{NS['w']}}}p")
    hyperlink = ET.SubElement(p, f"{{{NS['w']}}}hyperlink", {f"{{{NS['r']}}}id": rel_id})
    run = ET.SubElement(hyperlink, f"{{{NS['w']}}}r")
    run_pr = ET.SubElement(run, f"{{{NS['w']}}}rPr")
    ET.SubElement(run_pr, f"{{{NS['w']}}}rStyle", {f"{{{NS['w']}}}val": "Hyperlink"})
    text_el = ET.SubElement(run, f"{{{NS['w']}}}t")
    text_el.text = text
    return p


def create_paragraph_from_entry(entry: dict[str, Any], rels_root: ET.Element) -> ET.Element:
    if entry.get("kind") == "link":
        return create_hyperlink_paragraph(str(entry.get("text", "")), str(entry.get("target", "")), rels_root)
    return create_simple_paragraph(str(entry.get("text", "")))


def insert_paragraphs_after(body: ET.Element, anchor: ET.Element, entries: list[dict[str, Any]], rels_root: ET.Element) -> ET.Element:
    current = anchor
    for entry in entries:
        new_paragraph = create_paragraph_from_entry(entry, rels_root)
        idx = list(body).index(current)
        body.insert(idx + 1, new_paragraph)
        current = new_paragraph
    return current


def append_paragraphs_before_sectpr(body: ET.Element, entries: list[dict[str, Any]], rels_root: ET.Element) -> ET.Element | None:
    current: ET.Element | None = None
    insert_idx = len(body)
    for idx, child in enumerate(list(body)):
        if child.tag == f"{{{NS['w']}}}sectPr":
            insert_idx = idx
            break
    for entry in entries:
        new_paragraph = create_paragraph_from_entry(entry, rels_root)
        body.insert(insert_idx, new_paragraph)
        insert_idx += 1
        current = new_paragraph
    return current


def fill_appendix_sections(document_root: ET.Element, rels_root: ET.Element, context: dict[str, Any]) -> None:
    body = document_root.find("w:body", NS)
    if body is None:
        return

    sections = context.get("appendix_sections", []) or []
    if not sections:
        return

    direct_paragraphs = [child for child in list(body) if child.tag == f"{{{NS['w']}}}p"]
    paragraph_map = {paragraph_text(p): p for p in direct_paragraphs if paragraph_text(p)}

    last_inserted: ET.Element | None = None
    for section in sections:
        heading = str(section.get("heading", "")).strip()
        entries = [x for x in section.get("entries", []) if str(x.get("text", "")).strip()]
        if not heading or not entries:
            continue
        anchor = paragraph_map.get(heading)
        if anchor is None:
            created = append_paragraphs_before_sectpr(body, [_text_entry(heading), *entries], rels_root)
            if created is not None:
                last_inserted = created
            continue
        last_inserted = insert_paragraphs_after(body, anchor, entries, rels_root)


def render_pdf_to_jpeg_pages(pdf_path: Path, output_dir: Path, dpi: int = APPENDIX_RENDER_DPI) -> list[Path]:
    try:
        import fitz
    except ModuleNotFoundError as exc:
        raise ModuleNotFoundError("No module named 'fitz'。请安装 PyMuPDF 后再生成附件A/B。") from exc
    from PIL import Image

    out_paths: list[Path] = []
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    with fitz.open(pdf_path) as pdf:
        for page_idx, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            png_bytes = pix.tobytes("png")
            src = Image.open(io.BytesIO(png_bytes))
            img = src.convert("RGB")
            src.close()
            out_path = output_dir / f".appendix_render_{uuid.uuid4().hex}_{page_idx}.jpg"
            img.save(out_path, format="JPEG", quality=APPENDIX_JPEG_QUALITY, optimize=True)
            img.close()
            out_paths.append(out_path)
    return out_paths


def _normalize_appendix_file_entry(entry: Any) -> tuple[Path | None, str]:
    if isinstance(entry, dict):
        path_text = str(entry.get("path", "")).strip()
        caption = str(entry.get("caption", "")).strip()
        return (Path(path_text) if path_text else None), caption
    if isinstance(entry, (str, Path)):
        path_text = str(entry).strip()
        return (Path(path_text) if path_text else None), ""
    return None, ""


def insert_appendix_pdf_images(docx_path: Path, context: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    sections = []
    sections.extend(context.get("appendix_pdf_plan", []) or [])
    sections.extend(context.get("appendix_generated_plan", []) or [])
    if not sections:
        return

    def find_heading_paragraph(document: Any, heading: str) -> Any | None:
        for paragraph in document.paragraphs:
            if paragraph.text.strip() == heading:
                return paragraph
        return None

    def paragraph_is_blank(paragraph: Any) -> bool:
        if paragraph.text.strip():
            return False
        if paragraph._p.xpath(".//w:drawing"):
            return False
        return True

    def remove_paragraph(paragraph: Any) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    document = Document(docx_path)
    width = Inches(APPENDIX_IMAGE_WIDTH_INCHES)
    heading_set = {str(section.get("heading", "")).strip() for section in sections if str(section.get("heading", "")).strip()}

    def insert_paragraph_after(paragraph: Any, text: str = "") -> Any:
        new_para = document.add_paragraph()
        if text:
            new_para.add_run(text)
        paragraph._p.addnext(new_para._p)
        return new_para

    def insert_image_after(paragraph: Any, image_path: Path, caption: str = "") -> Any:
        current = paragraph
        if caption:
            caption_para = insert_paragraph_after(current, caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current = caption_para
        img_para = document.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(str(image_path), width=width)
        current._p.addnext(img_para._p)
        return img_para

    def clear_placeholder_paragraphs_after(anchor: Any) -> None:
        # The template leaves several blank paragraphs and page-break placeholders
        # after each appendix heading. Remove those placeholders first so the
        # appendix content starts from the first available blank spot on the
        # same page as the matched heading.
        paragraphs_snapshot = list(document.paragraphs)
        try:
            start_idx = paragraphs_snapshot.index(anchor) + 1
        except ValueError:
            return

        to_remove: list[Any] = []
        for paragraph in paragraphs_snapshot[start_idx:]:
            text = paragraph.text.strip()
            if text in heading_set:
                break
            if not paragraph_is_blank(paragraph):
                break
            to_remove.append(paragraph)

        for paragraph in to_remove:
            remove_paragraph(paragraph)

    temp_images: list[Path] = []
    try:
        for section in sections:
            heading = str(section.get("heading", "")).strip()
            files = [_normalize_appendix_file_entry(item) for item in section.get("files", [])]
            anchor = find_heading_paragraph(document, heading)
            if anchor is None or not files:
                continue

            clear_placeholder_paragraphs_after(anchor)
            current = anchor
            for file_path, caption in files:
                if file_path is None:
                    continue
                file_path = file_path.resolve()
                suffix = file_path.suffix.lower()
                if suffix in APPENDIX_IMAGE_SUFFIXES:
                    if file_path.exists():
                        current = insert_image_after(current, file_path, caption)
                    continue
                if suffix != ".pdf":
                    continue
                if caption:
                    current = insert_paragraph_after(current, caption)

                try:
                    image_paths = render_pdf_to_jpeg_pages(file_path, docx_path.parent)
                    temp_images.extend(image_paths)
                except Exception as exc:
                    current = insert_paragraph_after(current, f"PDF渲染失败：{file_path.name} ({type(exc).__name__}: {exc})")
                    continue

                for image_path in image_paths:
                    current = insert_image_after(current, image_path)
        document.save(docx_path)
    finally:
        for image_path in temp_images:
            try:
                image_path.unlink(missing_ok=True)
            except Exception:
                pass


def replace_dynamic_table_rows(table: ET.Element, rows_data: list[dict[str, Any]], header_rows: int, col_templates: list[str], env: Environment) -> None:
    rows = table.findall("w:tr", NS)
    if not rows:
        return

    # Older templates kept one sample data row after the headers and we cloned
    # that row as the body prototype. Once users clear the sample rows, the
    # table may only have header rows left. In that case, fall back to cloning
    # the last existing row so dynamic data can still be appended.
    prototype_idx = header_rows if len(rows) > header_rows else len(rows) - 1
    prototype = copy.deepcopy(rows[prototype_idx])

    if len(rows) > header_rows:
        for row_el in rows[header_rows:]:
            table.remove(row_el)

    if not rows_data:
        return

    compiled_templates = [env.from_string(tpl) for tpl in col_templates]
    for row_data in rows_data:
        new_row = copy.deepcopy(prototype)
        cells = new_row.findall("w:tc", NS)
        for idx, tpl in enumerate(compiled_templates):
            if idx >= len(cells):
                break
            val = tpl.render(row=row_data)
            set_cell_text(cells[idx], val)
        # Clear any remaining cells in the prototype row to avoid stale
        # template literals leaking into rendered data rows.
        for idx in range(len(compiled_templates), len(cells)):
            set_cell_text(cells[idx], "")
        table.append(new_row)


def set_table_cell(table: ET.Element, row_idx: int, col_idx: int, value: str) -> None:
    rows = table.findall("w:tr", NS)
    if row_idx >= len(rows):
        return
    cells = rows[row_idx].findall("w:tc", NS)
    if col_idx >= len(cells):
        return
    set_cell_text(cells[col_idx], value)


def fill_summary_tables(document_root: ET.Element, context: dict[str, Any]) -> None:
    # Table 9: member risk distribution
    table9 = find_table_by_locator(document_root, SUMMARY_TABLE_LOCATORS["member_risk_summary"])
    if table9 is not None:
        for idx, risk in enumerate(RISK_LEVEL_ORDER, start=1):
            set_table_cell(table9, 2, idx, str(context["member_risk_counts"].get(risk, 0)))
            set_table_cell(table9, 3, idx, context["member_risk_ratios"].get(risk, "0.00%"))

    # Table 10: node weld risk distribution by time (current + 5/10/15/20/25)
    table10 = find_table_by_locator(document_root, SUMMARY_TABLE_LOCATORS["node_risk_summary"])
    if table10 is not None:
        block_starts = [0, 4, 8, 12, 16, 20]
        for block_idx, block in enumerate(context["node_summary_blocks"][: len(block_starts)]):
            base = block_starts[block_idx]
            set_table_cell(table10, base, 0, block["title"])
            for ci, risk in enumerate(RISK_LEVEL_ORDER, start=1):
                set_table_cell(table10, base + 2, ci, str(block["counts"].get(risk, 0)))
                set_table_cell(table10, base + 3, ci, block["ratios"].get(risk, "0.00%"))

    # Table 13: member inspection summary
    table14 = find_table_by_locator(document_root, SUMMARY_TABLE_LOCATORS["member_inspection_summary"])
    if table14 is not None:
        for row_idx, risk in enumerate(RISK_LEVEL_ORDER, start=2):
            entry = next((x for x in context["member_inspection_summary"] if x["risk_level"] == risk), None)
            if entry is None:
                entry = {"count": 0, "II": 0, "III": 0, "IV": 0}
            set_table_cell(table14, row_idx, 0, risk)
            set_table_cell(table14, row_idx, 1, str(entry["count"]))
            set_table_cell(table14, row_idx, 2, "-" if entry["II"] == 0 else str(entry["II"]))
            set_table_cell(table14, row_idx, 3, "-" if entry["III"] == 0 else str(entry["III"]))
            set_table_cell(table14, row_idx, 4, "-" if entry["IV"] == 0 else str(entry["IV"]))
        set_table_cell(table14, 7, 1, str(context["member_inspection_total"]))
        set_table_cell(table14, 7, 2, str(context["member_inspection_total_II"]))
        set_table_cell(table14, 7, 3, str(context["member_inspection_total_III"]))
        set_table_cell(table14, 7, 4, str(context["member_inspection_total_IV"]))

    # Table 14: node inspection summary by future 5-year checkpoints
    table15 = find_table_by_locator(document_root, SUMMARY_TABLE_LOCATORS["node_inspection_summary"])
    if table15 is not None:
        block_starts_15 = [0, 8, 16, 24, 32]
        for block_idx, block in enumerate(context["node_inspection_blocks"][: len(block_starts_15)]):
            base = block_starts_15[block_idx]
            set_table_cell(table15, base, 0, block["time_node"])
            for ridx, risk in enumerate(RISK_LEVEL_ORDER, start=2):
                entry = next((x for x in block["summary_rows"] if x["risk_level"] == risk), None)
                if entry is None:
                    entry = {"count": 0, "II": 0, "III": 0, "IV": 0}
                set_table_cell(table15, ridx + base, 0, risk)
                set_table_cell(table15, ridx + base, 1, str(entry["count"]))
                set_table_cell(table15, ridx + base, 2, "-" if entry["II"] == 0 else str(entry["II"]))
                set_table_cell(table15, ridx + base, 3, "-" if entry["III"] == 0 else str(entry["III"]))
                set_table_cell(table15, ridx + base, 4, "-" if entry["IV"] == 0 else str(entry["IV"]))
            set_table_cell(table15, base + 7, 1, str(block["total_count"]))
            set_table_cell(table15, base + 7, 2, str(block["total_II"]))
            set_table_cell(table15, base + 7, 3, str(block["total_III"]))
            set_table_cell(table15, base + 7, 4, str(block["total_IV"]))


def should_render_detail_table(context: dict[str, Any], context_key: str) -> bool:
    if context_key not in OPTIONAL_WORD_DETAIL_CONTEXT_KEYS:
        return True
    return bool(context.get("include_word_plan_detail_tables", False))


def fill_cover_paragraphs(document_root: ET.Element, context: dict[str, Any], env: Environment) -> None:
    body = document_root.find("w:body", NS)
    if body is None:
        return

    cover_title_suffix = "\u5e73\u53f0\u98ce\u9669\u8bc4\u4f30\u53ca\u68c0\u6d4b\u7b56\u7565\u62a5\u544a"
    title_tpl = env.from_string("{{ platform_name }}" + cover_title_suffix)
    date_tpl = env.from_string("{{ report_date }}")

    for p in body.findall("w:p", NS):
        text = "".join((t.text or "") for t in p.findall(".//w:t", NS))
        if cover_title_suffix in text:
            set_paragraph_text(p, title_tpl.render(platform_name=context["platform_name"]))
            continue
        if "xx\u6708xx\u65e5" in text or "xxxx\u5e74xx\u6708xx\u65e5" in text:
            set_paragraph_text(p, date_tpl.render(report_date=context["report_date"]))


def render_text_placeholders(document_root: ET.Element, context: dict[str, Any], env: Environment) -> None:
    for paragraph in document_root.findall(".//w:p", NS):
        text = paragraph_text(paragraph)
        if not text or ("{{" not in text and "{%" not in text and "{#" not in text):
            continue
        try:
            rendered = env.from_string(text).render(context)
        except Exception as exc:
            detail = str(exc).strip()
            missing_match = re.search(r"'([^']+)' is undefined", detail)
            if missing_match:
                missing_key = missing_match.group(1)
                raise ValueError(f"Word 占位符渲染失败：缺少字段 `{missing_key}`。原文：{text}") from exc
            if detail:
                raise ValueError(f"Word 占位符渲染失败：{detail}。原文：{text}") from exc
            raise ValueError(f"Word 占位符渲染失败：{text}") from exc
        set_paragraph_text(paragraph, rendered)


def build_missing_requirements(context: dict[str, Any]) -> list[str]:
    missing: list[str] = []
    if to_text(context.get("platform_name")) == "":
        missing.append("封面平台名称(platform_name)缺失")
    if to_text(context.get("report_date")) == "":
        missing.append("封面报告日期(report_date)缺失")
    if len(context.get("fatigue_failure_rows", [])) == 0:
        missing.append("节点焊缝疲劳失效概率表无数据")
    if len(context.get("collapse_member_rows", [])) == 0:
        missing.append("构件倒塌失效概率表无数据")
    if len(context.get("collapse_joint_rows", [])) == 0:
        missing.append("节点倒塌失效概率表无数据")
    if len(context.get("node_risk_rows_current", [])) == 0:
        missing.append("当前节点风险等级表无数据")
    if len(context.get("member_risk_rows", [])) == 0:
        missing.append("构件风险等级表无数据")
    if context.get("include_word_plan_detail_tables", False) and len(context.get("member_inspection_rows", [])) == 0:
        missing.append("构件检验计划表无数据")
    if context.get("include_word_plan_detail_tables", False) and len(context.get("node_inspection_rows_future", [])) == 0:
        missing.append("节点检验计划(未来时间节点)表无数据")
    return missing


def build_row_cap_notes(context: dict[str, Any]) -> list[str]:
    total_counts = context.get("row_counts", {}) or {}
    limits = context.get("row_limits", {}) or {}

    mapping = {
        "fatigue_failure_rows": "fatigue_failure_rows_total",
        "collapse_member_rows": "collapse_member_rows_total",
        "collapse_joint_rows": "collapse_joint_rows_total",
        "node_risk_rows_current": "node_risk_rows_current_total",
        "member_risk_rows": "member_risk_rows_total",
        "member_inspection_rows": "member_inspection_rows_total",
        "node_inspection_rows_future": "node_inspection_rows_future_total",
    }

    notes: list[str] = []
    for key, total_key in mapping.items():
        shown = len(context.get(key, []))
        total = int(total_counts.get(total_key, shown))
        limit = limits.get(key)
        if isinstance(limit, int) and limit > 0 and total > shown:
            notes.append(f"{key}: {shown}/{total} (limited by {limit})")
    return notes


def render_report(template_docx: Path, output_docx: Path, context: dict[str, Any]) -> None:
    register_word_namespaces()
    env = Environment(undefined=StrictUndefined, autoescape=False, trim_blocks=True, lstrip_blocks=True)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    temp_output_docx = output_docx.with_name(f"{output_docx.stem}.{uuid.uuid4().hex}.tmp{output_docx.suffix}")

    with zipfile.ZipFile(template_docx, "r") as zin:
        doc_xml = zin.read("word/document.xml")
        root = ET.fromstring(doc_xml)
        rels_xml = zin.read("word/_rels/document.xml.rels")
        rels_root = ET.fromstring(rels_xml)
        try:
            settings_xml = zin.read("word/settings.xml")
            settings_root = ET.fromstring(settings_xml)
        except KeyError:
            settings_root = ET.Element(f"{{{NS['w']}}}settings")

        fill_cover_paragraphs(root, context, env)
        for spec in DYNAMIC_TABLE_SPECS:
            table = find_table_by_locator(root, spec.locator)
            if table is None:
                continue
            if not should_render_detail_table(context, spec.context_key):
                replace_dynamic_table_rows(
                    table=table,
                    rows_data=[],
                    header_rows=spec.header_rows,
                    col_templates=spec.column_templates,
                    env=env,
                )
                continue
            rows_data = context.get(spec.context_key, [])
            replace_dynamic_table_rows(
                table=table,
                rows_data=rows_data,
                header_rows=spec.header_rows,
                col_templates=spec.column_templates,
                env=env,
            )
        fill_summary_tables(root, context)
        if not context.get("appendix_pdf_plan") and not context.get("appendix_generated_plan"):
            fill_appendix_sections(root, rels_root, context)
        render_text_placeholders(root, context, env)
        rebuild_toc_field(root)
        mark_toc_fields_dirty(root)
        sanitize_markup_compatibility(root)
        sanitize_markup_compatibility(settings_root)
        ensure_update_fields_on_open(settings_root)

        out_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        out_rels_xml = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
        out_settings_xml = ET.tostring(settings_root, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(temp_output_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            wrote_settings = False
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, out_xml)
                elif item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item, out_rels_xml)
                elif item.filename == "word/settings.xml":
                    zout.writestr(item, out_settings_xml)
                    wrote_settings = True
                else:
                    zout.writestr(item, zin.read(item.filename))
            if not wrote_settings:
                zout.writestr("word/settings.xml", out_settings_xml)

    temp_output_docx.replace(output_docx)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render risk report docx using jinja2 and workbook data.")
    parser.add_argument("--workbook", required=True, help="Input xlsm workbook path.")
    parser.add_argument("--template", required=True, help="Input docx template path.")
    parser.add_argument("--output", required=True, help="Output rendered docx path.")
    parser.add_argument("--platform-name", default="", help="Platform display name for cover page.")
    parser.add_argument("--report-date", default="", help="Report date for cover page, e.g. 2026年3月12日.")
    parser.add_argument(
        "--metadata-json",
        default="",
        help="Optional metadata json file. Keys: platform_name, report_date. CLI args override file.",
    )
    parser.add_argument(
        "--full-rows",
        action="store_true",
        help="Deprecated: row caps are disabled by default.",
    )
    parser.add_argument(
        "--limit-rows",
        action="store_true",
        help="Enable row caps for large reports (disabled by default).",
    )
    parser.add_argument("--detail-row-limit", type=int, default=250, help="Row limit for Table 7/8/16 detail lists.")
    parser.add_argument("--future-row-limit", type=int, default=500, help="Total row limit for Table 17 future-node list.")
    parser.add_argument("--fatigue-row-limit", type=int, default=800, help="Row limit for Table 3 fatigue-failure list.")
    parser.add_argument("--collapse-row-limit", type=int, default=80, help="Row limit for Table 4/5 collapse-failure lists.")
    parser.add_argument("--appendix-a-file", default="", help="Appendix A file path.")
    parser.add_argument("--appendix-b-file", default="", help="Appendix B file path.")
    parser.add_argument("--appendix-c-dir", action="append", default=[], help="Appendix C directory path, repeatable.")
    parser.add_argument(
        "--include-word-plan-detail-tables",
        action="store_true",
        help="Fill Table 43/44 detail rows in Word. Default keeps these large detail tables only in Excel.",
    )
    parser.add_argument(
        "--non-vba-post-filters",
        action="store_true",
        help="Enable legacy non-VBA post filters/rebuilds in report context (default: strict VBA behavior).",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    workbook = Path(args.workbook).resolve()
    template = Path(args.template).resolve()
    output = Path(args.output).resolve()

    if not workbook.exists():
        raise FileNotFoundError(f"Workbook not found: {workbook}")
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")

    metadata: dict[str, Any] = {}
    if args.metadata_json:
        metadata_path = Path(args.metadata_json).resolve()
        if metadata_path.exists():
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))

    if args.platform_name:
        metadata["platform_name"] = args.platform_name
    if args.report_date:
        metadata["report_date"] = args.report_date

    row_limits: dict[str, int] | None = None
    if args.limit_rows and (not args.full_rows):
        row_limits = {
            "fatigue_failure_rows": args.fatigue_row_limit,
            "collapse_member_rows": args.collapse_row_limit,
            "collapse_joint_rows": args.collapse_row_limit,
            "node_risk_rows_current": args.detail_row_limit,
            "member_risk_rows": args.detail_row_limit,
            "member_inspection_rows": args.detail_row_limit,
            "node_inspection_rows_future": args.future_row_limit,
        }

    context = load_context_from_workbook(
        workbook,
        metadata,
        row_limits=row_limits,
        strict_vba_algorithms=not args.non_vba_post_filters,
    )
    context["appendix_sections"] = build_appendix_sections(
        appendix_a_file=args.appendix_a_file,
        appendix_b_file=args.appendix_b_file,
        appendix_c_dirs=list(args.appendix_c_dir),
    )
    context["appendix_pdf_plan"] = build_appendix_pdf_plan(
        appendix_a_file=args.appendix_a_file,
        appendix_b_file=args.appendix_b_file,
        appendix_c_dirs=list(args.appendix_c_dir),
    )
    context["include_word_plan_detail_tables"] = bool(args.include_word_plan_detail_tables)
    missing = build_missing_requirements(context)
    if missing:
        print("以下信息缺失或为空，可能影响模板完整输出：")
        for m in missing:
            print(f"- {m}")

    cap_notes = build_row_cap_notes(context)
    if cap_notes:
        print("已启用明细表限行（提升 Word 打开性能）：")
        for note in cap_notes:
            print(f"- {note}")

    render_report(template, output, context)
    if context.get("appendix_pdf_plan") or context.get("appendix_generated_plan"):
        insert_appendix_pdf_images(output, context)
    if not refresh_word_document_fields(output, pdf_output_path=output.with_suffix(".pdf")):
        raise RuntimeError(f"Word COM 自动更新目录并导出 PDF 失败：{output}")
    print(f"Report generated: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
