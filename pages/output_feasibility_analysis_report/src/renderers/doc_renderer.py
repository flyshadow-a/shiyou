from __future__ import annotations

from copy import deepcopy
import inspect
import re
from pathlib import Path
from typing import Any, Mapping, Sequence

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm
from docx.shared import RGBColor
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from src.config_loader import load_doc_renderer_config
from .table_writer import (
    find_table_by_header_row,
    find_tables_by_header_rows,
    write_basic_case_desc_table,
    write_basic_case_loads_table,
    write_combo_case_desc_table,
    write_combo_case_loads_table,
    write_analysis_summary_table,
    write_environment_marine_growth_table,
    write_environment_metric_table,
    write_environment_splash_zone_table,
    write_environment_water_level_table,
    write_pile_capacity_table,
    write_retrofit_history_table,
    write_cell,
)


DOC_RENDERER_CONFIG = load_doc_renderer_config()
TABLE_HEADERS = DOC_RENDERER_CONFIG["table_headers"]
PARAGRAPH_ANCHORS = DOC_RENDERER_CONFIG["paragraph_anchors"]
CHAPTER_PARAGRAPHS = DOC_RENDERER_CONFIG["chapter_paragraphs"]
DOC_STYLES = DOC_RENDERER_CONFIG["styles"]
RENDER_PLANS = DOC_RENDERER_CONFIG["render_plans"]

ANALYSIS_SUMMARY_HEADERS = TABLE_HEADERS["analysis_summary"]
PILE_CAPACITY_HEADER_ROWS = TABLE_HEADERS["pile_capacity"]
BASIC_CASE_DESC_HEADERS = TABLE_HEADERS["basic_case_desc"]
BASIC_CASE_LOADS_HEADERS = TABLE_HEADERS["basic_case_loads"]
COMBO_CASE_DESC_HEADERS = TABLE_HEADERS["combo_case_desc"]
COMBO_CASE_LOADS_HEADERS = TABLE_HEADERS["combo_case_loads"]
RETROFIT_HISTORY_LIST_HEADERS = TABLE_HEADERS["retrofit_history_list"]
RETROFIT_HISTORY_LIST_FALLBACK_HEADERS = ["序号", "项目名称", "年份"]
ENVIRONMENT_WATER_LEVEL_HEADERS = TABLE_HEADERS["environment_water_level"]
ENVIRONMENT_WAVE_HEADER_ROWS = TABLE_HEADERS["environment_wave"]
ENVIRONMENT_CURRENT_HEADER_ROWS = TABLE_HEADERS["environment_current"]
ENVIRONMENT_WIND_HEADER_ROWS = TABLE_HEADERS["environment_wind"]
ENVIRONMENT_MARINE_GROWTH_HEADER_ROWS = TABLE_HEADERS["environment_marine_growth"]
ENVIRONMENT_SPLASH_ZONE_HEADERS = TABLE_HEADERS["environment_splash_zone"]

MEMBER_SUMMARY_TITLE = PARAGRAPH_ANCHORS["member_summary_title"]
JOINT_SUMMARY_TITLE = PARAGRAPH_ANCHORS["joint_summary_title"]
PILE_SUMMARY_TITLE = PARAGRAPH_ANCHORS["pile_summary_title"]
RAW_BLOCK_PLACEHOLDER = PARAGRAPH_ANCHORS["raw_block_placeholder"]
PILE_EXTREME_LABEL = PARAGRAPH_ANCHORS["pile_extreme_label"]
PILE_OPERATION_LABEL = PARAGRAPH_ANCHORS["pile_operation_label"]
PILE_CAPACITY_CONCLUSION_PREFIX = PARAGRAPH_ANCHORS["pile_capacity_conclusion_prefix"]
PILE_STRESS_PREFIX = PARAGRAPH_ANCHORS["pile_stress_prefix"]
PILE_RAW_BLOCK_PREFIX = PARAGRAPH_ANCHORS["pile_raw_block_prefix"]
RAW_BLOCK_FONT_SIZE_PT = int(DOC_STYLES["raw_block_font_size_pt"])


def _normalize_text(text: str) -> str:
    return " ".join(text.split())


def _replace_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    font_size_pt: int | None = None,
    force_black: bool = False,
) -> None:
    element = paragraph._element
    template_rpr = None
    for run in paragraph.runs:
        r_pr = run._element.rPr
        if r_pr is not None:
            template_rpr = deepcopy(r_pr)
            break
    for child in list(element):
        if child.tag.endswith("}pPr"):
            continue
        element.remove(child)
    run = paragraph.add_run(text)
    if template_rpr is not None:
        run._element.insert(0, deepcopy(template_rpr))
    if force_black:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
        run.font.name = "Times New Roman"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), "宋体")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:cs"), "Times New Roman")
        half_points = str(font_size_pt * 2)
        sz = r_pr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            r_pr.append(sz)
        sz.set(qn("w:val"), half_points)
        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), half_points)


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _normalize_raw_block_line_for_paragraph(paragraph: Paragraph, text: str) -> str:
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return text.strip()
    return text.rstrip()


def _find_paragraph_index_by_exact_text(paragraphs: list[Paragraph], text: str) -> int:
    normalized_target = _normalize_text(text)
    for index, paragraph in enumerate(paragraphs):
        if _normalize_text(paragraph.text) == normalized_target:
            return index
    raise ValueError(f"未找到段落: {text}")


def _find_paragraph_index_by_prefix(paragraphs: list[Paragraph], prefix: str) -> int:
    normalized_prefix = _normalize_text(prefix)
    for index, paragraph in enumerate(paragraphs):
        if _normalize_text(paragraph.text).startswith(normalized_prefix):
            return index
    raise ValueError(f"未找到段落前缀: {prefix}")


def _find_paragraph_index_by_exact_text_after(
    paragraphs: list[Paragraph], text: str, start_index: int
) -> int:
    normalized_target = _normalize_text(text)
    for index in range(start_index + 1, len(paragraphs)):
        if _normalize_text(paragraphs[index].text) == normalized_target:
            return index
    raise ValueError(f"未找到索引 {start_index} 之后的段落: {text}")


def _find_next_nonempty_paragraph_index(paragraphs: list[Paragraph], start_index: int) -> int:
    for index in range(start_index + 1, len(paragraphs)):
        if _normalize_text(paragraphs[index].text):
            return index
    raise ValueError(f"未找到索引 {start_index} 之后的正文段落")


def _find_next_paragraph_index(paragraphs: list[Paragraph], start_index: int) -> int:
    next_index = start_index + 1
    if next_index >= len(paragraphs):
        raise ValueError(f"未找到索引 {start_index} 之后的段落")
    return next_index


def _find_next_title_index(paragraphs: list[Paragraph], titles: Sequence[str], start_index: int) -> int:
    # 1～3 章渲染现在按“章节区间”工作，这里用于找到当前标题之后的下一个章节标题，
    # 从而把替换范围限定在本章节正文内部。
    normalized_titles = {_normalize_text(title) for title in titles if title.strip()}
    for index in range(start_index + 1, len(paragraphs)):
        if _normalize_text(paragraphs[index].text) in normalized_titles:
            return index
    return len(paragraphs)


def _find_last_nonempty_paragraph(paragraphs: Sequence[Paragraph]) -> Paragraph | None:
    for paragraph in reversed(paragraphs):
        if _normalize_text(paragraph.text):
            return paragraph
    return None


def _delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _clone_paragraph_properties(source: Paragraph, target: Paragraph) -> None:
    source_ppr = source._element.pPr
    if source_ppr is None:
        return

    target_ppr = target._element.pPr
    if target_ppr is not None:
        target._element.remove(target_ppr)
    target._element.insert(0, deepcopy(source_ppr))


def _insert_paragraph_after(paragraph: Paragraph, text: str, *, copy_from: Paragraph | None = None) -> Paragraph:
    new_paragraph = OxmlElement("w:p")
    paragraph._element.addnext(new_paragraph)
    inserted = Paragraph(new_paragraph, paragraph._parent)
    if copy_from is not None:
        if copy_from.style is not None:
            inserted.style = copy_from.style
        _clone_paragraph_properties(copy_from, inserted)
    elif paragraph.style is not None:
        inserted.style = paragraph.style
    _replace_paragraph_text(inserted, text)
    return inserted


def _replace_paragraph_region(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
    lines: Sequence[str],
    font_size_pt: int | None = None,
) -> None:
    paragraphs = list(document.paragraphs)
    region = paragraphs[start_index:end_index]
    if not region:
        raise ValueError(f"待替换段落区间为空: {start_index}-{end_index}")

    normalized_lines = [line.rstrip() for line in lines]
    for index, line in enumerate(normalized_lines[: len(region)]):
        _replace_paragraph_text(
            region[index],
            _normalize_raw_block_line_for_paragraph(region[index], line),
            font_size_pt=font_size_pt,
        )

    template_paragraph = region[-1]
    last_paragraph = region[min(len(region), len(normalized_lines)) - 1] if normalized_lines else region[0]
    for line in normalized_lines[len(region) :]:
        normalized_line = _normalize_raw_block_line_for_paragraph(template_paragraph, line)
        last_paragraph = _insert_paragraph_after(
            last_paragraph, normalized_line, copy_from=template_paragraph
        )
        if font_size_pt is not None:
            _replace_paragraph_text(last_paragraph, normalized_line, font_size_pt=font_size_pt)

    for paragraph in reversed(region[len(normalized_lines) :]):
        _delete_paragraph(paragraph)


def _resolve_region_template_paragraph(
    region: Sequence[Paragraph],
    *,
    style_source: str,
) -> Paragraph:
    if style_source == "last_nonempty":
        return _find_last_nonempty_paragraph(region) or region[-1]
    return region[0]


def _rewrite_paragraph_region(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
    lines: Sequence[str],
    template_paragraph: Paragraph,
    font_size_pt: int | None = None,
) -> None:
    paragraphs = list(document.paragraphs)
    region = paragraphs[start_index:end_index]
    if not region:
        raise ValueError(f"待重写段落区间为空: {start_index}-{end_index}")

    anchor_paragraph = paragraphs[start_index - 1]
    for paragraph in reversed(region):
        _delete_paragraph(paragraph)

    last_paragraph = anchor_paragraph
    for line in lines:
        normalized_line = _normalize_raw_block_line_for_paragraph(template_paragraph, line)
        last_paragraph = _insert_paragraph_after(
            last_paragraph,
            normalized_line,
            copy_from=template_paragraph,
        )
        if font_size_pt is not None:
            _replace_paragraph_text(last_paragraph, normalized_line, font_size_pt=font_size_pt)


def _find_paragraph_in_region_by_prefix(region: Sequence[Paragraph], prefix: str) -> int:
    normalized_prefix = _normalize_text(prefix)
    for index, paragraph in enumerate(region):
        if _normalize_text(paragraph.text).startswith(normalized_prefix):
            return index
    raise ValueError(f"未找到章节示例段落前缀: {prefix}")


def _find_paragraph_in_region_by_prefix_occurrence(
    region: Sequence[Paragraph],
    prefix: str,
    occurrence: int,
) -> int:
    # 1.1 平台概述里会出现多个“例子：...”段落，
    # 这里用 occurrence 精确定位第几个示例段，避免第二段误替换第一段。
    normalized_prefix = _normalize_text(prefix)
    target_occurrence = max(1, int(occurrence))
    current_occurrence = 0
    for index, paragraph in enumerate(region):
        if _normalize_text(paragraph.text).startswith(normalized_prefix):
            current_occurrence += 1
            if current_occurrence == target_occurrence:
                return index
    raise ValueError(f"未找到章节示例段落前缀: {prefix} (第 {target_occurrence} 个)")


def _replace_paragraph_at_anchor(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
    anchor_prefix: str,
    lines: Sequence[str],
) -> None:
    paragraphs = list(document.paragraphs)
    region = paragraphs[start_index:end_index]
    if not region:
        raise ValueError(f"待替换章节区间为空: {start_index}-{end_index}")

    anchor_offset = _find_paragraph_in_region_by_prefix(region, anchor_prefix)
    anchor_paragraph = region[anchor_offset]
    style_paragraph = None
    for paragraph in region[anchor_offset + 1 :]:
        if _normalize_text(paragraph.text):
            style_paragraph = paragraph
            break
    if style_paragraph is None:
        style_paragraph = anchor_paragraph

    normalized_lines = [
        _normalize_raw_block_line_for_paragraph(style_paragraph, line)
        for line in lines
        if line.strip()
    ]
    if not normalized_lines:
        return

    if style_paragraph.style is not None:
        anchor_paragraph.style = style_paragraph.style
    _clone_paragraph_properties(style_paragraph, anchor_paragraph)
    _replace_paragraph_text(anchor_paragraph, normalized_lines[0], force_black=True)

    last_paragraph = anchor_paragraph
    for line in normalized_lines[1:]:
        last_paragraph = _insert_paragraph_after(last_paragraph, line, copy_from=style_paragraph)
        _replace_paragraph_text(last_paragraph, line, force_black=True)


def _replace_paragraph_blocks_at_anchors(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
    blocks: Sequence[Mapping[str, Any]],
) -> bool:
    # 当前专门给 1.1 平台概述使用：每个 block 自带锚点信息。
    # 除了直接匹配“例子：...”段落外，也支持先命中说明段，再替换其后的下一段正文。
    paragraphs = list(document.paragraphs)
    region = paragraphs[start_index:end_index]
    if not region:
        raise ValueError(f"待替换章节区间为空: {start_index}-{end_index}")

    # 先基于原始章节区间锁定所有待替换锚点，
    # 避免前一个 block 已经替换后，后一个 block 的 occurrence 序号发生偏移。
    anchor_targets: list[tuple[Paragraph, Paragraph, str, Paragraph | None, bool]] = []
    for block in blocks:
        if not isinstance(block, Mapping):
            continue
        line = str(block.get("text", "")).strip()
        anchor_prefix = str(block.get("anchor_prefix", "")).strip()
        occurrence = int(block.get("anchor_occurrence", 1) or 1)
        preserve_anchor_style = bool(block.get("preserve_anchor_style", False))
        replace_next_paragraph = bool(block.get("replace_next_paragraph", False))
        keep_anchor_paragraph = bool(block.get("keep_anchor_paragraph", False))
        if not line or not anchor_prefix:
            continue

        anchor_offset = _find_paragraph_in_region_by_prefix_occurrence(
            region,
            anchor_prefix,
            occurrence,
        )
        marker_paragraph = region[anchor_offset]
        anchor_paragraph = marker_paragraph
        if replace_next_paragraph:
            # 某些模板先给一行括号说明，再给下一段“例子：...”正文；
            # 这种情况下需要用说明行定位，但最终输出里不再保留该提示行。
            for paragraph in region[anchor_offset + 1 :]:
                if _normalize_text(paragraph.text):
                    anchor_paragraph = paragraph
                    break
        if preserve_anchor_style:
            style_paragraph = anchor_paragraph
        else:
            style_paragraph = None
            for paragraph in region[anchor_offset + 1 :]:
                if _normalize_text(paragraph.text):
                    style_paragraph = paragraph
                    break
            if style_paragraph is None:
                style_paragraph = anchor_paragraph
        anchor_targets.append(
            (
                anchor_paragraph,
                style_paragraph,
                line,
                marker_paragraph if replace_next_paragraph else None,
                keep_anchor_paragraph,
            )
        )

    replaced_any = False
    for anchor_paragraph, style_paragraph, line, marker_paragraph, keep_anchor_paragraph in anchor_targets:
        if style_paragraph.style is not None:
            anchor_paragraph.style = style_paragraph.style
        _clone_paragraph_properties(style_paragraph, anchor_paragraph)
        _replace_paragraph_text(
            anchor_paragraph,
            _normalize_raw_block_line_for_paragraph(style_paragraph, line),
            force_black=True,
        )
        if (
            marker_paragraph is not None
            and marker_paragraph is not anchor_paragraph
            and not keep_anchor_paragraph
        ):
            _delete_paragraph(marker_paragraph)
        replaced_any = True

    return replaced_any


def _replace_placeholders_in_region(
    paragraphs: Sequence[Paragraph],
    placeholders: Mapping[str, Any],
) -> None:
    # hybrid / preserve_template 模式下保留模板正文结构，只替换显式占位符。
    normalized_placeholders = {
        f"{{{{{key}}}}}": str(value).strip()
        for key, value in placeholders.items()
        if str(key).strip()
    }
    if not normalized_placeholders:
        return

    for paragraph in paragraphs:
        text = paragraph.text
        replaced = text
        for token, value in normalized_placeholders.items():
            replaced = replaced.replace(token, value)
        if replaced != text:
            _replace_paragraph_text(paragraph, replaced)


def _clear_platform_overview_template_prompts(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
) -> None:
    paragraphs = list(document.paragraphs)[start_index:end_index]
    for paragraph in paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        if text.startswith("例子：") or re.match(r"^（第.+部分：", text):
            _delete_paragraph(paragraph)


def _clear_retrofit_history_template_prompts(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
) -> None:
    paragraphs = list(document.paragraphs)[start_index:end_index]
    for paragraph in paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        if re.match(r"^（来自历次改造信息里.*列表.*）$", text):
            _delete_paragraph(paragraph)


def _clear_platform_evaluation_conclusion_template_prompts(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
) -> None:
    paragraphs = list(document.paragraphs)[start_index:end_index]
    for paragraph in paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        if text == "（来自本次改造的信息和评估结论）" or text.startswith("例子："):
            _delete_paragraph(paragraph)


def _clear_environment_conditions_template_prompts(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
) -> None:
    prompt_texts = {
        "（来自平台所在油气田信息的水深水位信息）",
        "（来自平台所在油气田信息的波浪信息）",
        "（来自平台所在油气田信息的海流信息）",
        "（来自平台所在油气田信息的风信息）",
        "（来自平台基本信息的海生物信息）",
        "（来自平台基本信息的飞溅区腐蚀余量）",
        "海生物信息如下所示。（来自平台基本信息的海生物信息）",
        "海生物信息如下所示。（来自平台基本信息的飞溅区腐蚀余量）",
        "（该部分内容内置）",
        "（来自文件管理相应平台的文件）",
    }
    paragraphs = list(document.paragraphs)[start_index:end_index]
    for paragraph in paragraphs:
        text = _normalize_text(paragraph.text)
        if text in prompt_texts:
            _delete_paragraph(paragraph)


def _clear_basis_data_template_prompts(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
) -> None:
    paragraphs = list(document.paragraphs)[start_index:end_index]
    for paragraph in paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        if text == "（来自文件管理相应平台的文件）" or text == "……" or re.match(r"^（\d+）xxx$", text):
            _delete_paragraph(paragraph)


def _clear_load_information_template_prompts(
    document: DocxDocument,
    *,
    start_index: int,
    end_index: int,
) -> None:
    paragraphs = list(document.paragraphs)[start_index:end_index]
    for paragraph in paragraphs:
        text = _normalize_text(paragraph.text)
        if text == "（来自载荷信息相应平台的载荷信息表，表格一致）":
            _delete_paragraph(paragraph)


def _clear_fixed_second_chapter_prompts(document: DocxDocument) -> None:
    section_prompts = {
        "规范依据": {"（该部分内容内置）"},
        "评估流程": {"（该部分内容内置）"},
        "基础数据": {"（来自文件管理相应平台的文件）"},
    }
    titles = [title for title in section_prompts]
    trailing_titles = titles[1:] + ["环境条件"]

    for index, title in enumerate(titles):
        paragraphs = list(document.paragraphs)
        try:
            title_index = _find_paragraph_index_by_exact_text(paragraphs, title)
        except ValueError:
            continue
        start_index = _find_next_paragraph_index(paragraphs, title_index)
        end_index = _find_next_title_index(paragraphs, [trailing_titles[index]], title_index)
        for paragraph in list(document.paragraphs)[start_index:end_index]:
            if _normalize_text(paragraph.text) in section_prompts[title]:
                _delete_paragraph(paragraph)


def _remove_paragraph_text_fragment(paragraph: Paragraph, fragment: str) -> None:
    if fragment not in paragraph.text:
        return
    _replace_paragraph_text(paragraph, paragraph.text.replace(fragment, "").rstrip())


def _clear_fixed_analysis_model_prompts(document: DocxDocument) -> None:
    paragraphs = list(document.paragraphs)
    try:
        section_start = _find_paragraph_index_by_exact_text(paragraphs, "程序和坐标系统")
    except ValueError:
        return

    try:
        section_end = _find_paragraph_index_by_exact_text_after(paragraphs, "设计水平分析", section_start)
    except ValueError:
        section_end = _find_next_title_index(paragraphs, ["基本工况", "构件名义应力校核"], section_start)

    figure_prompt_fragments = (
        "（根据模型节点构件信息绘制）",
        "（根据模型节点构件及水平层信息绘制）",
    )
    prompt_texts = {"（该部分内容内置）"}
    for paragraph in list(document.paragraphs)[section_start + 1:section_end]:
        text = _normalize_text(paragraph.text)
        if text in prompt_texts:
            _delete_paragraph(paragraph)
            continue
        for fragment in figure_prompt_fragments:
            _remove_paragraph_text_fragment(paragraph, fragment)


def _clear_design_level_analysis_template_prompts(document: DocxDocument) -> None:
    paragraphs = list(document.paragraphs)
    try:
        section_start = _find_paragraph_index_by_exact_text(paragraphs, "设计水平分析")
        section_end = _find_paragraph_index_by_exact_text_after(paragraphs, "构件名义应力校核", section_start)
    except ValueError:
        return

    delete_paragraph_texts = {
        "（以上内容内置）",
    }
    prompt_fragments = (
        "（该部分内容内置）",
        "（内容内置）",
        "（来自结果读取的基本工况信息）",
        "（来自结果读取的组合工况信息）",
        "（内置参考结论性语句）",
        "（来自结果读取结果的汇总信息）",
        "（该内容数据和理论来自模型/海况文件，读取代码后续开发）",
    )
    for paragraph in list(document.paragraphs)[section_start + 1:section_end]:
        text = _normalize_text(paragraph.text)
        if text in delete_paragraph_texts:
            _delete_paragraph(paragraph)
            continue
        for fragment in prompt_fragments:
            _remove_paragraph_text_fragment(paragraph, fragment)


def _clear_45_template_prompts(document: DocxDocument) -> None:
    paragraphs = list(document.paragraphs)
    try:
        section_start = _find_paragraph_index_by_exact_text(paragraphs, MEMBER_SUMMARY_TITLE)
    except ValueError:
        return

    delete_paragraph_texts = {
        "（以下内容来自结果文件读取）",
    }
    prompt_fragments = (
        "（校核规范名称通过结果文件读取可获得）",
        "（根据结果更新最小安全系数）",
        "（更新最大UC值，以下内容来自结果文件读取）",
    )
    for paragraph in list(document.paragraphs)[section_start + 1:]:
        text = _normalize_text(paragraph.text)
        if text in delete_paragraph_texts:
            _delete_paragraph(paragraph)
            continue
        for fragment in prompt_fragments:
            _remove_paragraph_text_fragment(paragraph, fragment)


def _insert_picture_before_paragraph(
    paragraph: Paragraph,
    image_path: str,
    *,
    width_cm: float = 14.0,
) -> None:
    image = Path(str(image_path or "").strip())
    if not image.exists() or not image.is_file():
        return

    picture_element = OxmlElement("w:p")
    paragraph._element.addprevious(picture_element)
    picture_paragraph = Paragraph(picture_element, paragraph._parent)
    if paragraph.style is not None:
        picture_paragraph.style = paragraph.style
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image), width=Cm(width_cm))


def _replace_paragraph_with_picture(
    paragraph: Paragraph,
    image_path: str,
    *,
    width_cm: float = 14.0,
) -> None:
    image = Path(str(image_path or "").strip())
    if not image.exists() or not image.is_file():
        return

    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("}pPr"):
            continue
        element.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Cm(width_cm))


def _replace_or_insert_picture_before_paragraph(
    document: DocxDocument,
    paragraph_index: int,
    image_path: str,
    *,
    width_cm: float = 14.0,
) -> None:
    paragraphs = list(document.paragraphs)
    target_paragraph = paragraphs[paragraph_index]
    if paragraph_index > 0:
        previous_paragraph = paragraphs[paragraph_index - 1]
        if previous_paragraph._element.xpath(".//w:drawing"):
            _replace_paragraph_with_picture(previous_paragraph, image_path, width_cm=width_cm)
            return
    _insert_picture_before_paragraph(target_paragraph, image_path, width_cm=width_cm)


def _insert_analysis_model_images(document: DocxDocument, section: Mapping[str, Any]) -> None:
    overall_model_image_path = str(section.get("overall_model_image_path", "")).strip()
    coordinate_system_image_path = str(section.get("coordinate_system_image_path", "")).strip()
    paragraphs = list(document.paragraphs)
    if overall_model_image_path:
        try:
            caption_index = _find_paragraph_index_by_prefix(paragraphs, "图3.1 整体模型")
        except ValueError:
            caption_index = -1
        if caption_index >= 0:
            _replace_or_insert_picture_before_paragraph(
                document,
                caption_index,
                overall_model_image_path,
            )
    if coordinate_system_image_path:
        paragraphs = list(document.paragraphs)
        try:
            caption_index = _find_paragraph_index_by_prefix(paragraphs, "图3.2")
        except ValueError:
            caption_index = -1
        if caption_index >= 0:
            _replace_or_insert_picture_before_paragraph(
                document,
                caption_index,
                coordinate_system_image_path,
            )


def _normalize_chapter_section(section: Any) -> dict[str, Any]:
    # 渲染层再次做一次兜底归一化，确保 report_service 或测试即使继续传旧格式字符串，
    # 也能按新的章节块级逻辑稳定处理。
    if isinstance(section, Mapping):
        blocks = section.get("blocks")
        normalized_blocks = []
        if isinstance(blocks, Sequence) and not isinstance(blocks, (str, bytes)):
            for block in blocks:
                if not isinstance(block, Mapping):
                    continue
                text = str(block.get("text", "")).strip()
                if not text:
                    continue
                normalized_blocks.append(
                    {
                        "kind": str(block.get("kind", "generated_paragraph")).strip()
                        or "generated_paragraph",
                        "text": text,
                        # 为章节内多示例段替换预留锚点元数据；
                        # 普通 block 没有这些字段时会自动忽略。
                        "anchor_prefix": str(block.get("anchor_prefix", "")).strip(),
                        "anchor_occurrence": int(block.get("anchor_occurrence", 1) or 1),
                        # 某些示例段本身就是目标正文格式，
                        # 可通过该开关要求直接保留锚点段自身格式。
                        "preserve_anchor_style": bool(block.get("preserve_anchor_style", False)),
                        # 某些 block 需要替换“锚点后的下一段”，而不是锚点本身。
                        "replace_next_paragraph": bool(block.get("replace_next_paragraph", False)),
                        "keep_anchor_paragraph": bool(block.get("keep_anchor_paragraph", False)),
                    }
                )

        if not normalized_blocks:
            text = str(section.get("text", "")).strip()
            if text:
                normalized_blocks = [{"kind": "generated_paragraph", "text": text}]

        placeholders = section.get("placeholders", {})
        if not isinstance(placeholders, Mapping):
            placeholders = {}

        return {
            "mode": str(section.get("mode", "replace_region")).strip() or "replace_region",
            "blocks": normalized_blocks,
            "placeholders": placeholders,
            "table_rows": list(section.get("table_rows", [])),
            "overall_model_image_path": str(section.get("overall_model_image_path", "")).strip(),
            "coordinate_system_image_path": str(section.get("coordinate_system_image_path", "")).strip(),
            "load_information_meta": dict(section.get("load_information_meta", {})),
            "load_information_rows": list(section.get("load_information_rows", [])),
            "water_level_rows": list(section.get("water_level_rows", [])),
            "wind_rows": list(section.get("wind_rows", [])),
            "wave_rows": list(section.get("wave_rows", [])),
            "current_rows": list(section.get("current_rows", [])),
            "marine_growth_rows": list(section.get("marine_growth_rows", [])),
            "splash_zone_rows": list(section.get("splash_zone_rows", [])),
        }

    text = str(section or "").strip()
    return {
        "mode": "replace_region",
        "blocks": ([{"kind": "generated_paragraph", "text": text}] if text else []),
        "placeholders": {},
        "table_rows": [],
        "overall_model_image_path": "",
        "coordinate_system_image_path": "",
        "load_information_meta": {},
        "load_information_rows": [],
        "water_level_rows": [],
        "wind_rows": [],
        "wave_rows": [],
        "current_rows": [],
        "marine_growth_rows": [],
        "splash_zone_rows": [],
    }


def _group_environment_metric_rows(rows: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], dict[str, Any]] = {}
    ordered_keys: list[tuple[str, str]] = []
    for row in rows:
        if not isinstance(row, Mapping):
            continue
        group_name = str(row.get("group_name", "")).strip()
        item_name = str(row.get("item_name", "")).strip()
        return_period = str(row.get("return_period", "")).strip()
        key = (group_name, item_name)
        if key not in grouped:
            grouped[key] = {
                "group_name": group_name,
                "item_name": item_name,
                "values_by_period": {},
            }
            ordered_keys.append(key)
        grouped[key]["values_by_period"][return_period] = str(row.get("value", "")).strip()
    return [grouped[key] for key in ordered_keys]


def render_environment_conditions_tables(
    document: DocxDocument,
    *,
    chapter_1_3_context: Mapping[str, Any] | None = None,
) -> None:
    if not chapter_1_3_context:
        return

    environment_section = chapter_1_3_context.get("environment_conditions", {})
    if not isinstance(environment_section, Mapping):
        return

    water_level_rows = environment_section.get("water_level_rows", [])
    wave_rows = environment_section.get("wave_rows", [])
    current_rows = environment_section.get("current_rows", [])
    wind_rows = environment_section.get("wind_rows", [])
    marine_growth_rows = environment_section.get("marine_growth_rows", [])
    splash_zone_rows = environment_section.get("splash_zone_rows", [])
    has_environment_rows = any(
        isinstance(rows, Sequence) and not isinstance(rows, (str, bytes)) and rows
        for rows in (water_level_rows, wave_rows, current_rows, wind_rows, marine_growth_rows, splash_zone_rows)
    )
    if not has_environment_rows:
        return

    if isinstance(water_level_rows, Sequence) and not isinstance(water_level_rows, (str, bytes)) and water_level_rows:
        water_level_table = find_table_by_header_row(document.tables, ENVIRONMENT_WATER_LEVEL_HEADERS)
        write_environment_water_level_table(water_level_table, water_level_rows)

    if isinstance(wave_rows, Sequence) and not isinstance(wave_rows, (str, bytes)) and wave_rows:
        wave_table_matches = find_tables_by_header_rows(document.tables, ENVIRONMENT_WAVE_HEADER_ROWS)
        if not wave_table_matches:
            raise ValueError("未找到 2.5 节波浪参数表")
        write_environment_metric_table(wave_table_matches[0], _group_environment_metric_rows(wave_rows))

    if isinstance(current_rows, Sequence) and not isinstance(current_rows, (str, bytes)) and current_rows:
        current_table_matches = find_tables_by_header_rows(document.tables, ENVIRONMENT_CURRENT_HEADER_ROWS)
        if not current_table_matches:
            raise ValueError("未找到 2.5 节海流参数表")
        write_environment_metric_table(current_table_matches[0], _group_environment_metric_rows(current_rows))

    if isinstance(wind_rows, Sequence) and not isinstance(wind_rows, (str, bytes)) and wind_rows:
        wind_table_matches = find_tables_by_header_rows(document.tables, ENVIRONMENT_WIND_HEADER_ROWS)
        if not wind_table_matches:
            raise ValueError("未找到 2.5 节风参数表")
        write_environment_metric_table(wind_table_matches[0], _group_environment_metric_rows(wind_rows))

    if isinstance(marine_growth_rows, Sequence) and not isinstance(marine_growth_rows, (str, bytes)) and marine_growth_rows:
        marine_growth_table_matches = find_tables_by_header_rows(document.tables, ENVIRONMENT_MARINE_GROWTH_HEADER_ROWS)
        if not marine_growth_table_matches:
            raise ValueError("未找到 2.5.5 节海生物信息表")
        write_environment_marine_growth_table(marine_growth_table_matches[0], marine_growth_rows)

    if isinstance(splash_zone_rows, Sequence) and not isinstance(splash_zone_rows, (str, bytes)) and splash_zone_rows:
        splash_zone_table = find_table_by_header_row(document.tables, ENVIRONMENT_SPLASH_ZONE_HEADERS)
        write_environment_splash_zone_table(splash_zone_table, splash_zone_rows)


def _render_chapter_section(
    document: DocxDocument,
    *,
    config: Mapping[str, Any],
    section: Mapping[str, Any],
    all_titles: Sequence[str],
) -> None:
    # 新的 1～3 章策略：
    # - 先根据 title/end_title 锁定章节正文范围
    # - 再按 mode 决定是保留模板、整段替换还是混合插入
    paragraphs = list(document.paragraphs)
    title = str(config.get("title", "")).strip()
    if not title:
        return
    is_platform_overview = title == "平台概况"
    is_retrofit_history = title == "改造历史"
    is_platform_evaluation_conclusion = title == "平台的评估结论"
    is_basis_data = title == "基础数据"
    is_load_information = title == "载荷变化"
    is_environment_conditions = title == "环境条件"
    is_analysis_model = title == "分析模型"

    title_index = _find_paragraph_index_by_exact_text(paragraphs, title)
    start_index = _find_next_paragraph_index(paragraphs, title_index)
    explicit_end_title = str(config.get("end_title", "")).strip()
    end_titles = [explicit_end_title] if explicit_end_title else list(all_titles)
    end_index = _find_next_title_index(paragraphs, end_titles, title_index)
    if explicit_end_title and end_index == len(paragraphs):
        end_index = _find_next_title_index(paragraphs, list(all_titles), title_index)
    if start_index > end_index:
        return

    region = paragraphs[start_index:end_index]
    placeholders = section.get("placeholders", {})
    if isinstance(placeholders, Mapping):
        _replace_placeholders_in_region(region, placeholders)
    if is_analysis_model:
        _insert_analysis_model_images(document, section)

    mode = str(section.get("mode") or config.get("body_mode") or "replace_region").strip()
    if mode == "preserve_template":
        return

    blocks = section.get("blocks", [])
    if not isinstance(blocks, Sequence) or isinstance(blocks, (str, bytes)):
        blocks = []
    lines = [str(block.get("text", "")).strip() for block in blocks if isinstance(block, Mapping)]
    lines = [line for line in lines if line]
    if not lines:
        if is_retrofit_history:
            _clear_retrofit_history_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_load_information:
            _clear_load_information_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_environment_conditions:
            _clear_environment_conditions_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        return

    style_source = str(config.get("replace_style_source", "first")).strip() or "first"
    rewrite_region = str(config.get("rewrite_region", "false")).strip().lower() == "true"
    replace_anchor_prefix = str(config.get("replace_anchor_prefix", "")).strip()

    if mode == "hybrid":
        # hybrid 模式保留原正文，并在正文区最后一个有效段落后追加动态段落。
        # 这样新增内容会继承示例正文样式，而不是标题或“例子”提示行的样式。
        insert_anchor = str(config.get("insert_anchor", "")).strip()
        current_paragraphs = list(document.paragraphs)
        if insert_anchor:
            anchor_index = _find_paragraph_index_by_exact_text_after(
                current_paragraphs, insert_anchor, title_index
            )
            anchor_paragraph = current_paragraphs[anchor_index]
            template_paragraph = current_paragraphs[
                _find_next_nonempty_paragraph_index(current_paragraphs, anchor_index)
            ]
        else:
            anchor_paragraph = _find_last_nonempty_paragraph(region)
            template_paragraph = anchor_paragraph

        if anchor_paragraph is None or template_paragraph is None:
            anchor_paragraph = current_paragraphs[start_index - 1]
            template_paragraph = current_paragraphs[start_index] if start_index < end_index else anchor_paragraph

        last_paragraph = anchor_paragraph
        for line in lines:
            last_paragraph = _insert_paragraph_after(last_paragraph, line, copy_from=template_paragraph)
        if is_platform_overview:
            _clear_platform_overview_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_retrofit_history:
            _clear_retrofit_history_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_platform_evaluation_conclusion:
            _clear_platform_evaluation_conclusion_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_basis_data:
            _clear_basis_data_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_load_information:
            _clear_load_information_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_environment_conditions:
            _clear_environment_conditions_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        return

    # 如果 block 自带锚点信息，则优先按 block 粒度定点替换。
    # 这用于 1.1 节内存在两个“例子：...”段落的情况。
    if _replace_paragraph_blocks_at_anchors(
        document,
        start_index=start_index,
        end_index=end_index,
        blocks=[block for block in blocks if isinstance(block, Mapping)],
    ):
        if is_platform_overview:
            _clear_platform_overview_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_retrofit_history:
            _clear_retrofit_history_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_basis_data:
            _clear_basis_data_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_load_information:
            _clear_load_information_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_platform_evaluation_conclusion:
            _clear_platform_evaluation_conclusion_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_environment_conditions:
            _clear_environment_conditions_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        return

    if replace_anchor_prefix:
        try:
            _replace_paragraph_at_anchor(
                document,
                start_index=start_index,
                end_index=end_index,
                anchor_prefix=replace_anchor_prefix,
                lines=lines,
            )
            if is_platform_overview:
                _clear_platform_overview_template_prompts(
                    document,
                    start_index=start_index,
                    end_index=end_index,
                )
            if is_retrofit_history:
                _clear_retrofit_history_template_prompts(
                    document,
                    start_index=start_index,
                    end_index=end_index,
                )
            if is_basis_data:
                _clear_basis_data_template_prompts(
                    document,
                    start_index=start_index,
                    end_index=end_index,
                )
            if is_load_information:
                _clear_load_information_template_prompts(
                    document,
                    start_index=start_index,
                    end_index=end_index,
                )
            if is_platform_evaluation_conclusion:
                _clear_platform_evaluation_conclusion_template_prompts(
                    document,
                    start_index=start_index,
                    end_index=end_index,
                )
            if is_environment_conditions:
                _clear_environment_conditions_template_prompts(
                    document,
                    start_index=start_index,
                    end_index=end_index,
                )
            return
        except ValueError:
            pass

    if rewrite_region:
        template_paragraph = _resolve_region_template_paragraph(region, style_source=style_source)
        _rewrite_paragraph_region(
            document,
            start_index=start_index,
            end_index=end_index,
            lines=lines,
            template_paragraph=template_paragraph,
        )
        if is_platform_overview:
            _clear_platform_overview_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_retrofit_history:
            _clear_retrofit_history_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_basis_data:
            _clear_basis_data_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_load_information:
            _clear_load_information_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_platform_evaluation_conclusion:
            _clear_platform_evaluation_conclusion_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        if is_environment_conditions:
            _clear_environment_conditions_template_prompts(
                document,
                start_index=start_index,
                end_index=end_index,
            )
        return

    _replace_paragraph_region(
        document,
        start_index=start_index,
        end_index=end_index,
        lines=lines,
    )
    if is_platform_overview:
        _clear_platform_overview_template_prompts(
            document,
            start_index=start_index,
            end_index=end_index,
        )
    if is_retrofit_history:
        _clear_retrofit_history_template_prompts(
            document,
            start_index=start_index,
            end_index=end_index,
        )
    if is_basis_data:
        _clear_basis_data_template_prompts(
            document,
            start_index=start_index,
            end_index=end_index,
        )
    if is_load_information:
        _clear_load_information_template_prompts(
            document,
            start_index=start_index,
            end_index=end_index,
        )
    if is_platform_evaluation_conclusion:
        _clear_platform_evaluation_conclusion_template_prompts(
            document,
            start_index=start_index,
            end_index=end_index,
        )
    if is_environment_conditions:
        _clear_environment_conditions_template_prompts(
            document,
            start_index=start_index,
            end_index=end_index,
        )


def _build_pile_capacity_conclusion(
    *,
    extreme_compression: Mapping[str, Any],
    extreme_tension: Mapping[str, Any],
    operation_compression: Mapping[str, Any],
    operation_tension: Mapping[str, Any],
) -> str:
    def summarize_condition(name: str, summaries: list[Mapping[str, Any]]) -> str:
        valid_items = [item for item in summaries if item.get("is_pass_text") != "无数据"]
        if not valid_items:
            return f"{name}相关结果暂无数据，暂无法判断是否满足规范要求。"

        min_sf = min(float(item.get("min_sf", 0.0)) for item in valid_items)
        all_pass = all(str(item.get("is_pass_text", "")) == "满足" for item in valid_items)
        status_text = "满足规范要求" if all_pass else "不满足规范要求"
        return f"{name}平台桩承载力控制最小安全系数为{min_sf:.2f}，{status_text}。"

    extreme_text = summarize_condition("极端工况下", [extreme_compression, extreme_tension])
    operation_text = summarize_condition("操作工况下", [operation_compression, operation_tension])
    return f"从以上结果可知：{extreme_text}{operation_text}"


def _load_document(template_path: str) -> DocxDocument:
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"模板不存在: {template}")
    return Document(str(template))


def _save_document(document: DocxDocument, output_path: str) -> str:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return str(output)


def _copy_element_relationships(source_document: DocxDocument, target_document: DocxDocument, element: Any) -> None:
    relationship_attributes = (qn("r:embed"), qn("r:id"), qn("r:link"))
    for node in element.iter():
        for attribute in relationship_attributes:
            old_rel_id = node.get(attribute)
            if not old_rel_id or old_rel_id not in source_document.part.rels:
                continue
            source_rel = source_document.part.rels[old_rel_id]
            if source_rel.is_external:
                new_rel_id = target_document.part.relate_to(
                    source_rel.target_ref,
                    source_rel.reltype,
                    is_external=True,
                )
            else:
                new_rel_id = target_document.part.relate_to(
                    source_rel.target_part,
                    source_rel.reltype,
                )
            node.set(attribute, new_rel_id)


def _append_copied_body_elements(
    target_document: DocxDocument,
    source_document: DocxDocument,
    elements: Sequence[Any],
) -> None:
    target_body = target_document.element.body
    section_properties = target_body.sectPr
    for source_element in elements:
        copied_element = deepcopy(source_element)
        _copy_element_relationships(source_document, target_document, copied_element)
        if section_properties is None:
            target_body.append(copied_element)
        else:
            target_body.insert(target_body.index(section_properties), copied_element)


def _find_appendix_a_elements(source_document: DocxDocument) -> list[Any]:
    def is_appendix_title(text: str, appendix_name: str) -> bool:
        return text.startswith(f"{appendix_name}：") or text.startswith(f"{appendix_name}:")

    paragraphs = list(source_document.paragraphs)
    appendix_a_indexes = [
        index
        for index, paragraph in enumerate(paragraphs)
        if is_appendix_title(_normalize_text(paragraph.text), "附录A")
    ]
    if not appendix_a_indexes:
        raise ValueError("参考报告中未找到附录A")

    start_paragraph = paragraphs[appendix_a_indexes[-1]]
    end_paragraph: Paragraph | None = None
    for paragraph in paragraphs[appendix_a_indexes[-1] + 1:]:
        text = _normalize_text(paragraph.text)
        if is_appendix_title(text, "附录B") or is_appendix_title(text, "附录C"):
            end_paragraph = paragraph
            break

    body_elements = list(source_document.element.body)
    start_index = body_elements.index(start_paragraph._element)
    end_index = body_elements.index(end_paragraph._element) if end_paragraph is not None else len(body_elements)
    return [element for element in body_elements[start_index:end_index] if element.tag != qn("w:sectPr")]


def append_appendix_a_from_reference(
    document: DocxDocument,
    *,
    reference_path: str,
) -> None:
    reference = Path(reference_path)
    if not reference.exists() or not reference.is_file():
        return
    reference_document = Document(str(reference))
    appendix_elements = _find_appendix_a_elements(reference_document)
    if not appendix_elements:
        return
    document.add_page_break()
    _append_copied_body_elements(document, reference_document, appendix_elements)


def _find_appendix_a_title_paragraph(document: DocxDocument) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if _normalize_text(paragraph.text).startswith("附录A"):
            return paragraph
    return None


def _copy_first_run_properties(source: Paragraph, target: Paragraph) -> None:
    if not source.runs or not target.runs:
        return
    source_rpr = source.runs[0]._element.rPr
    if source_rpr is not None:
        target.runs[0]._element.insert(0, deepcopy(source_rpr))


def _append_appendix_title(document: DocxDocument, title_text: str) -> Paragraph:
    title = document.add_paragraph(title_text)
    appendix_a_title = _find_appendix_a_title_paragraph(document)
    if appendix_a_title is not None and appendix_a_title is not title:
        if appendix_a_title.style is not None:
            title.style = appendix_a_title.style
        _clone_paragraph_properties(appendix_a_title, title)
        _copy_first_run_properties(appendix_a_title, title)
    else:
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.bold = True
    return title


def append_empty_appendix_b(document: DocxDocument) -> None:
    document.add_page_break()
    _append_appendix_title(document, "附录B：新增井槽改造方案")


def _normalize_marker_text(text: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "", str(text or "").upper())


def _find_line_index_by_marker(lines: Sequence[str], marker: str, *, start_index: int = 0) -> int:
    normalized_marker = _normalize_marker_text(marker)
    for index in range(max(0, start_index), len(lines)):
        if normalized_marker in _normalize_marker_text(lines[index]):
            return index
    raise ValueError(f"未找到原始块标记: {marker}")


def _extract_appendix_c_pile_lines(lines: Sequence[str]) -> list[str]:
    start_index = _find_line_index_by_marker(lines, "PSIOPT")
    return [line.rstrip("\n\r") for line in lines[start_index:]]


def _append_raw_text_lines(
    document: DocxDocument,
    lines: Sequence[str],
    *,
    font_size_pt: int = RAW_BLOCK_FONT_SIZE_PT,
    font_name: str = "Courier New",
) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(font_size_pt + 1)
        run = paragraph.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)


def append_appendix_c_from_factor(
    document: DocxDocument,
    *,
    factor_lines: Sequence[str],
) -> None:
    missing_message = "未能找到完整 PSI 输入桩基数据，请检查当前平台运行目录下 psiinp* 文件是否存在且内容完整。"
    try:
        appendix_lines = _extract_appendix_c_pile_lines(factor_lines)
    except ValueError:
        appendix_lines = [missing_message]
    if not appendix_lines:
        appendix_lines = [missing_message]
    document.add_page_break()
    _append_appendix_title(document, "附录C：桩基数据")
    _append_raw_text_lines(document, appendix_lines, font_size_pt=7.5, font_name="宋体_GB2312")


def _resolve_default_appendix_a_reference(template_path: str) -> str:
    return str(Path(template_path).with_name("xxx平台改建可行性评估报告.docx"))


def _render_cover_fields(
    document: DocxDocument,
    *,
    cover_platform_name: str = "",
    report_date_text: str = "",
) -> None:
    paragraphs = list(document.paragraphs)
    normalized_platform_name = str(cover_platform_name or "").strip()
    if normalized_platform_name:
        title_index = _find_paragraph_index_by_prefix(paragraphs, "XXX平台结构强度改造可行性评估报告")
        _replace_paragraph_text(
            paragraphs[title_index],
            f"{normalized_platform_name}平台结构强度改造可行性评估报告",
        )
        header_title = f"{normalized_platform_name}平台结构强度改造可行性评估"
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                if "XXX平台结构强度改造可行性评估" in paragraph.text:
                    _replace_paragraph_text(
                        paragraph,
                        paragraph.text.replace("XXX平台结构强度改造可行性评估", header_title),
                    )
    if report_date_text:
        date_index = _find_paragraph_index_by_exact_text(paragraphs, "xxxx年xx月xx日")
        _replace_paragraph_text(paragraphs[date_index], report_date_text)


def _find_pile_capacity_tables(document: DocxDocument) -> list[Any]:
    pile_tables = find_tables_by_header_rows(document.tables, PILE_CAPACITY_HEADER_ROWS)
    if len(pile_tables) < 2:
        raise ValueError("未找到 4.5.3 的两张桩基承载力表")
    return pile_tables


def render_43_tables(
    document: DocxDocument,
    *,
    basic_case_desc_rows: Sequence[Mapping[str, Any]],
    basic_case_load_rows: Sequence[Mapping[str, Any]],
) -> None:
    basic_case_desc_table = find_table_by_header_row(document.tables, BASIC_CASE_DESC_HEADERS)
    basic_case_loads_table = find_table_by_header_row(document.tables, BASIC_CASE_LOADS_HEADERS)

    write_basic_case_desc_table(basic_case_desc_table, basic_case_desc_rows)
    write_basic_case_loads_table(basic_case_loads_table, basic_case_load_rows)


def render_44_tables(
    document: DocxDocument,
    *,
    combo_case_desc_rows: Sequence[Mapping[str, Any]],
    combo_case_load_rows: Sequence[Mapping[str, Any]],
) -> None:
    combo_case_desc_table = find_table_by_header_row(document.tables, COMBO_CASE_DESC_HEADERS)
    combo_case_loads_table = find_table_by_header_row(document.tables, COMBO_CASE_LOADS_HEADERS)

    write_combo_case_desc_table(combo_case_desc_table, combo_case_desc_rows)
    write_combo_case_loads_table(combo_case_loads_table, combo_case_load_rows)


def render_analysis_summary_table_section(
    document: DocxDocument,
    *,
    analysis_summary: Mapping[str, Any],
) -> None:
    summary_table = find_table_by_header_row(document.tables, ANALYSIS_SUMMARY_HEADERS)
    write_analysis_summary_table(summary_table, list(analysis_summary.get("items", [])))


def render_45_3_capacity_tables(
    document: DocxDocument,
    *,
    pile_axial_capacity_summary: Mapping[str, Any],
) -> None:
    pile_tables = _find_pile_capacity_tables(document)
    write_pile_capacity_table(
        pile_tables[0],
        list(pile_axial_capacity_summary.get("extreme_table_rows", [])),
    )
    write_pile_capacity_table(
        pile_tables[1],
        list(pile_axial_capacity_summary.get("operation_table_rows", [])),
    )


def render_45_summary_section(
    document: DocxDocument,
    *,
    member_summary: Mapping[str, Any] | None = None,
    joint_summary: Mapping[str, Any] | None = None,
    pile_stress_summary: Mapping[str, Any] | None = None,
    pile_axial_capacity_summary: Mapping[str, Any] | None = None,
) -> None:
    write_45_summary_paragraphs(
        document,
        member_summary=member_summary,
        joint_summary=joint_summary,
        pile_stress_summary=pile_stress_summary,
        pile_axial_capacity_summary=pile_axial_capacity_summary,
    )


def render_45_raw_block_section(
    document: DocxDocument,
    *,
    member_group_summary: Mapping[str, Any] | None = None,
    joint_can_summary_result: Mapping[str, Any] | None = None,
    pile_group_summary_result: Mapping[str, Any] | None = None,
) -> None:
    write_45_raw_blocks(
        document,
        member_group_summary=member_group_summary,
        joint_can_summary_result=joint_can_summary_result,
        pile_group_summary_result=pile_group_summary_result,
    )


def render_1_3_chapter_paragraphs(
    document: DocxDocument,
    *,
    chapter_1_3_context: Mapping[str, Any] | None = None,
) -> None:
    if not chapter_1_3_context:
        return

    # 先收集所有章节标题，供每个章节计算自己的正文结束边界。
    all_titles = [
        str(config.get("title", "")).strip()
        for config in CHAPTER_PARAGRAPHS.values()
        if str(config.get("title", "")).strip()
    ]
    for key, config in CHAPTER_PARAGRAPHS.items():
        section = _normalize_chapter_section(chapter_1_3_context.get(key, {}))
        has_table_rows = bool(section.get("table_rows"))
        has_environment_rows = any(
            bool(section.get(table_key))
            for table_key in (
                "water_level_rows",
                "wind_rows",
                "wave_rows",
                "current_rows",
                "marine_growth_rows",
                "splash_zone_rows",
            )
        )
        has_load_information = bool(section.get("load_information_meta")) or bool(
            section.get("load_information_rows")
        )
        has_analysis_model_images = bool(section.get("overall_model_image_path")) or bool(
            section.get("coordinate_system_image_path")
        )
        if (
            not section.get("blocks")
            and not section.get("placeholders")
            and not has_table_rows
            and not has_load_information
            and not has_environment_rows
            and not has_analysis_model_images
        ):
            continue
        _render_chapter_section(
            document,
            config=config,
            section=section,
            all_titles=all_titles,
        )


def render_retrofit_history_list_table(
    document: DocxDocument,
    *,
    chapter_1_3_context: Mapping[str, Any] | None = None,
) -> None:
    if not chapter_1_3_context:
        return

    retrofit_section = chapter_1_3_context.get("retrofit_history", {})
    if not isinstance(retrofit_section, Mapping):
        return

    table_rows = retrofit_section.get("table_rows", [])
    if not isinstance(table_rows, Sequence) or isinstance(table_rows, (str, bytes)) or not table_rows:
        return

    # 1.2 节模板存在两种表头版本：
    # - 旧版：序号 / 改造项目 / 年份
    # - 当前纯净模板：序号 / 项目名称 / 年份
    # 这里优先匹配业务语义更清晰的旧版表头，找不到时再回退到当前模板表头。
    try:
        retrofit_table = find_table_by_header_row(document.tables, RETROFIT_HISTORY_LIST_HEADERS)
    except ValueError:
        retrofit_table = find_table_by_header_row(
            document.tables,
            RETROFIT_HISTORY_LIST_FALLBACK_HEADERS,
        )
    write_retrofit_history_table(retrofit_table, table_rows)


def _move_table_after_paragraph(document: DocxDocument, paragraph: Paragraph, rows: int, cols: int):
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table_element = table._tbl
    table_element.getparent().remove(table_element)
    paragraph._element.addnext(table_element)
    return table


def _set_load_information_cell_margins(cell, margin_twips: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def _reset_load_information_paragraph_character_indents(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)

    for attr_name in ("firstLineChars", "leftChars", "rightChars", "hangingChars"):
        ind.set(qn(f"w:{attr_name}"), "0")


def _write_load_information_cell(
    cell,
    value: str,
    *,
    font_size_pt: int = 7,
    bold: bool = False,
) -> None:
    write_cell(cell, value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_load_information_cell_margins(cell)
    for paragraph in cell.paragraphs:
        _reset_load_information_paragraph_character_indents(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(font_size_pt + 1)
        for run in paragraph.runs:
            run.font.size = Pt(font_size_pt)
            run.bold = bold


def _set_load_information_column_widths(table) -> None:
    table.autofit = False
    widths = [
        0.45,
        1.25,
        0.85,
        1.45,
        0.85,
        0.85,
        0.75,
        1.05,
        0.75,
        0.55,
        0.55,
        0.55,
        0.55,
        0.55,
        0.55,
        0.70,
        0.70,
        0.55,
        1.00,
    ]
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Cm(width)


def _set_load_information_meta_row(table, row_index: int, definitions: Sequence[tuple[int, int, str, str]]) -> None:
    for start, total_span, label, value in definitions:
        label_span = 3 if total_span == 7 else 2
        value_span = total_span - label_span
        label_cell = table.cell(row_index, start)
        if label_span > 1:
            label_cell = label_cell.merge(table.cell(row_index, start + label_span - 1))
        _write_load_information_cell(label_cell, label, bold=True)

        value_cell = table.cell(row_index, start + label_span)
        if value_span > 1:
            value_cell = value_cell.merge(table.cell(row_index, start + total_span - 1))
        _write_load_information_cell(value_cell, value)


def _build_load_information_table(
    document: DocxDocument,
    *,
    anchor_paragraph: Paragraph,
    load_information_meta: Mapping[str, Any],
    load_information_rows: Sequence[Mapping[str, Any]],
) -> None:
    table = _move_table_after_paragraph(
        document,
        anchor_paragraph,
        rows=4 + len(load_information_rows),
        cols=19,
    )
    _set_load_information_column_widths(table)

    _set_load_information_meta_row(
        table,
        0,
        (
            (0, 7, "所属分公司", str(load_information_meta.get("branch", ""))),
            (7, 6, "所属作业单元", str(load_information_meta.get("op_company", ""))),
            (13, 6, "所属油（气）田", str(load_information_meta.get("oilfield", ""))),
        ),
    )
    _set_load_information_meta_row(
        table,
        1,
        (
            (0, 7, "设施名称", str(load_information_meta.get("facility_name", ""))),
            (7, 6, "投产时间", str(load_information_meta.get("start_time", ""))),
            (13, 6, "设计年限", str(load_information_meta.get("design_life", ""))),
        ),
    )

    merged_row_headers = {
        0: "序号",
        1: "改扩建\n项目名称",
        2: "改扩建\n时间",
        3: "改扩建\n内容",
        4: "上部组块\n总操作重量\n（MT）",
        5: "上部组块\n不可超越重量\n（MT）",
        6: "重量变化\n（MT）",
        7: "上部组块重心\nx,y,z\n（m）",
        8: "重心不可\n超越半径\n（m）",
        17: "是否整体\n评估",
        18: "评估机构",
    }
    for column_index, text in merged_row_headers.items():
        merged_cell = table.cell(2, column_index).merge(table.cell(3, column_index))
        _write_load_information_cell(merged_cell, text, bold=True)

    _write_load_information_cell(
        table.cell(2, 9).merge(table.cell(2, 14)),
        "极端工况最大载荷",
        bold=True,
    )
    _write_load_information_cell(
        table.cell(2, 15).merge(table.cell(2, 16)),
        "桩基承载力\n安全系数（最小）",
        bold=True,
    )

    sub_headers = [
        "Fx\n（KN）",
        "Fy\n（KN）",
        "Fz\n（KN）",
        "Mx\n（KN·m）",
        "My\n（KN·m）",
        "Mz,（KN·m）",
        "操作工况",
        "极端工况",
    ]
    for column_index, text in enumerate(sub_headers, start=9):
        _write_load_information_cell(table.cell(3, column_index), text, bold=True)

    row_fields = [
        "seq_no",
        "project_name",
        "rebuild_time",
        "rebuild_content",
        "total_weight_mt",
        "weight_limit_mt",
        "weight_delta_mt",
        "center_xyz",
        "center_radius_m",
        "fx_kn",
        "fy_kn",
        "fz_kn",
        "mx_kn_m",
        "my_kn_m",
        "mz_kn_m",
        "safety_op",
        "safety_extreme",
        "overall_assessment",
        "assessment_org",
    ]
    for row_index, item in enumerate(load_information_rows, start=4):
        for column_index, field in enumerate(row_fields):
            _write_load_information_cell(table.cell(row_index, column_index), str(item.get(field, "")))


def render_load_information_table_section(
    document: DocxDocument,
    *,
    chapter_1_3_context: Mapping[str, Any] | None = None,
) -> None:
    if not chapter_1_3_context:
        return

    load_information_section = chapter_1_3_context.get("load_information", {})
    if not isinstance(load_information_section, Mapping):
        return

    load_information_meta = load_information_section.get("load_information_meta", {})
    load_information_rows = load_information_section.get("load_information_rows", [])
    if not isinstance(load_information_meta, Mapping):
        return
    if not isinstance(load_information_rows, Sequence) or isinstance(load_information_rows, (str, bytes)):
        return
    if not load_information_meta and not load_information_rows:
        return

    paragraphs = list(document.paragraphs)
    try:
        anchor_index = _find_paragraph_index_by_exact_text(paragraphs, "荷载情况如下表：")
    except ValueError:
        anchor_index = _find_paragraph_index_by_exact_text(paragraphs, "载荷变化")
    insert_anchor = _insert_paragraph_after(paragraphs[anchor_index], "", copy_from=paragraphs[anchor_index])
    _build_load_information_table(
        document,
        anchor_paragraph=insert_anchor,
        load_information_meta=load_information_meta,
        load_information_rows=load_information_rows,
    )


SECTION_RENDERERS = {
    "chapters_1_3": render_1_3_chapter_paragraphs,
    "retrofit_history_list_table": render_retrofit_history_list_table,
    "load_information_table": render_load_information_table_section,
    "environment_conditions_tables": render_environment_conditions_tables,
    "tables_43": render_43_tables,
    "tables_44": render_44_tables,
    "analysis_summary_table": render_analysis_summary_table_section,
    "pile_capacity_tables": render_45_3_capacity_tables,
    "summary_paragraphs_45": render_45_summary_section,
    "raw_blocks_45": render_45_raw_block_section,
}


def _render_plan(document: DocxDocument, plan_name: str, **context: Any) -> None:
    sections = RENDER_PLANS.get(plan_name, [])
    for section_name in sections:
        renderer = SECTION_RENDERERS.get(section_name)
        if renderer is None:
            raise ValueError(f"未注册的渲染章节: {section_name}")
        parameter_names = set(inspect.signature(renderer).parameters)
        section_kwargs = {
            key: value for key, value in context.items() if key in parameter_names
        }
        renderer(document, **section_kwargs)


def write_45_summary_paragraphs(
    document: DocxDocument,
    *,
    member_summary: Mapping[str, Any] | None = None,
    joint_summary: Mapping[str, Any] | None = None,
    pile_stress_summary: Mapping[str, Any] | None = None,
    pile_axial_capacity_summary: Mapping[str, Any] | None = None,
) -> None:
    paragraphs = list(document.paragraphs)

    if member_summary is not None:
        title_index = _find_paragraph_index_by_exact_text(paragraphs, MEMBER_SUMMARY_TITLE)
        body_index = _find_next_nonempty_paragraph_index(paragraphs, title_index)
        _replace_paragraph_text(paragraphs[body_index], str(member_summary.get("summary_text", "")))

    if joint_summary is not None:
        title_index = _find_paragraph_index_by_exact_text(paragraphs, JOINT_SUMMARY_TITLE)
        body_index = _find_next_nonempty_paragraph_index(paragraphs, title_index)
        _replace_paragraph_text(paragraphs[body_index], str(joint_summary.get("summary_text", "")))

    if pile_axial_capacity_summary is not None:
        extreme_label_index = _find_paragraph_index_by_exact_text(paragraphs, PILE_EXTREME_LABEL)
        operation_label_index = _find_paragraph_index_by_exact_text(paragraphs, PILE_OPERATION_LABEL)
        conclusion_index = _find_paragraph_index_by_prefix(
            paragraphs, PILE_CAPACITY_CONCLUSION_PREFIX
        )
        conclusion_text = _build_pile_capacity_conclusion(
            extreme_compression=pile_axial_capacity_summary.get("extreme_compression", {}),
            extreme_tension=pile_axial_capacity_summary.get("extreme_tension", {}),
            operation_compression=pile_axial_capacity_summary.get("operation_compression", {}),
            operation_tension=pile_axial_capacity_summary.get("operation_tension", {}),
        )

        _replace_paragraph_text(paragraphs[_find_next_paragraph_index(paragraphs, extreme_label_index)], "")
        _replace_paragraph_text(paragraphs[_find_next_paragraph_index(paragraphs, operation_label_index)], "")
        _replace_paragraph_text(paragraphs[conclusion_index], conclusion_text)

    if pile_stress_summary is not None:
        pile_stress_index = _find_paragraph_index_by_prefix(paragraphs, PILE_STRESS_PREFIX)
        _replace_paragraph_text(paragraphs[pile_stress_index], str(pile_stress_summary.get("summary_text", "")))


def write_45_raw_blocks(
    document: DocxDocument,
    *,
    member_group_summary: Mapping[str, Any] | None = None,
    joint_can_summary_result: Mapping[str, Any] | None = None,
    pile_group_summary_result: Mapping[str, Any] | None = None,
) -> None:
    paragraphs = list(document.paragraphs)

    if member_group_summary is not None:
        member_title_index = _find_paragraph_index_by_exact_text(paragraphs, MEMBER_SUMMARY_TITLE)
        member_placeholder_index = _find_paragraph_index_by_exact_text_after(
            paragraphs, RAW_BLOCK_PLACEHOLDER, member_title_index
        )
        joint_title_index = _find_paragraph_index_by_exact_text(paragraphs, JOINT_SUMMARY_TITLE)
        _replace_paragraph_region(
            document,
            start_index=_find_next_paragraph_index(paragraphs, member_placeholder_index),
            end_index=joint_title_index,
            lines=str(member_group_summary.get("raw_block", "")).splitlines(),
            font_size_pt=RAW_BLOCK_FONT_SIZE_PT,
        )
        paragraphs = list(document.paragraphs)

    if joint_can_summary_result is not None:
        joint_title_index = _find_paragraph_index_by_exact_text(paragraphs, JOINT_SUMMARY_TITLE)
        joint_placeholder_index = _find_paragraph_index_by_exact_text_after(
            paragraphs, RAW_BLOCK_PLACEHOLDER, joint_title_index
        )
        pile_title_index = _find_paragraph_index_by_exact_text(paragraphs, PILE_SUMMARY_TITLE)
        _replace_paragraph_region(
            document,
            start_index=_find_next_paragraph_index(paragraphs, joint_placeholder_index),
            end_index=pile_title_index,
            lines=str(joint_can_summary_result.get("raw_block", "")).splitlines(),
            font_size_pt=RAW_BLOCK_FONT_SIZE_PT,
        )
        paragraphs = list(document.paragraphs)

    if pile_group_summary_result is not None:
        pile_raw_start_index = _find_paragraph_index_by_prefix(paragraphs, PILE_RAW_BLOCK_PREFIX)
        _replace_paragraph_region(
            document,
            start_index=pile_raw_start_index,
            end_index=len(paragraphs),
            lines=str(pile_group_summary_result.get("raw_block", "")).splitlines(),
            font_size_pt=RAW_BLOCK_FONT_SIZE_PT,
        )


def render_analysis_summary_doc(
    *,
    template_path: str,
    output_path: str,
    analysis_summary: Mapping[str, Any],
) -> str:
    document = _load_document(template_path)
    _render_plan(document, "analysis_summary_doc", analysis_summary=analysis_summary)
    return _save_document(document, output_path)


def render_45_3_tables(
    *,
    template_path: str,
    output_path: str,
    analysis_summary: Mapping[str, Any],
    pile_axial_capacity_summary: Mapping[str, Any],
    member_summary: Mapping[str, Any] | None = None,
    joint_summary: Mapping[str, Any] | None = None,
    pile_stress_summary: Mapping[str, Any] | None = None,
) -> str:
    document = _load_document(template_path)

    _render_plan(
        document,
        "section_45_3_doc",
        analysis_summary=analysis_summary,
        pile_axial_capacity_summary=pile_axial_capacity_summary,
        member_summary=member_summary,
        joint_summary=joint_summary,
        pile_stress_summary=pile_stress_summary,
    )

    return _save_document(document, output_path)


def render_report_doc(
    *,
    template_path: str,
    output_path: str,
    appendix_a_reference_path: str | None = None,
    appendix_c_factor_lines: Sequence[str] | None = None,
    cover_platform_name: str = "",
    report_date_text: str = "",
    analysis_summary: Mapping[str, Any],
    pile_axial_capacity_summary: Mapping[str, Any],
    basic_case_desc_rows: Sequence[Mapping[str, Any]],
    basic_case_load_rows: Sequence[Mapping[str, Any]],
    combo_case_desc_rows: Sequence[Mapping[str, Any]],
    combo_case_load_rows: Sequence[Mapping[str, Any]],
    chapter_1_3_context: Mapping[str, Any] | None = None,
    member_summary: Mapping[str, Any] | None = None,
    joint_summary: Mapping[str, Any] | None = None,
    pile_stress_summary: Mapping[str, Any] | None = None,
    member_group_summary: Mapping[str, Any] | None = None,
    joint_can_summary_result: Mapping[str, Any] | None = None,
    pile_group_summary_result: Mapping[str, Any] | None = None,
) -> str:
    document = _load_document(template_path)

    _render_cover_fields(
        document,
        cover_platform_name=cover_platform_name,
        report_date_text=report_date_text,
    )

    _render_plan(
        document,
        "report_doc",
        basic_case_desc_rows=basic_case_desc_rows,
        basic_case_load_rows=basic_case_load_rows,
        combo_case_desc_rows=combo_case_desc_rows,
        combo_case_load_rows=combo_case_load_rows,
        chapter_1_3_context=chapter_1_3_context,
        analysis_summary=analysis_summary,
        pile_axial_capacity_summary=pile_axial_capacity_summary,
        member_summary=member_summary,
        joint_summary=joint_summary,
        pile_stress_summary=pile_stress_summary,
        member_group_summary=member_group_summary,
        joint_can_summary_result=joint_can_summary_result,
        pile_group_summary_result=pile_group_summary_result,
    )

    _clear_fixed_second_chapter_prompts(document)
    _clear_fixed_analysis_model_prompts(document)
    _clear_design_level_analysis_template_prompts(document)
    _clear_45_template_prompts(document)
    append_appendix_a_from_reference(
        document,
        reference_path=appendix_a_reference_path or _resolve_default_appendix_a_reference(template_path),
    )
    append_empty_appendix_b(document)
    if appendix_c_factor_lines is not None:
        append_appendix_c_from_factor(document, factor_lines=appendix_c_factor_lines)

    return _save_document(document, output_path)
