from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_MODULE_ROOT = PROJECT_ROOT / "pages" / "output_feasibility_analysis_report"
if str(REPORT_MODULE_ROOT) not in sys.path:
    sys.path.insert(0, str(REPORT_MODULE_ROOT))

from src.renderers.table_writer import (  # noqa: E402
    NON_BREAKING_HYPHEN,
    find_table_by_header_row,
    write_analysis_summary_table,
    write_basic_case_loads_table,
    write_combo_case_desc_table,
    write_environment_marine_growth_table,
    write_pile_capacity_table,
)


class ReportTableWriterTests(unittest.TestCase):
    def test_find_table_by_header_row_accepts_legacy_splash_zone_unit(self) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=3)
        headers = ["飞溅区上限(m)", "飞溅区下限(m)", "腐蚀余量(mm/y)"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        found = find_table_by_header_row(
            document.tables,
            ["飞溅区上限(m)", "飞溅区下限(m)", "腐蚀余量(mm)"],
            alternate_headers=[headers],
        )

        self.assertEqual("腐蚀余量(mm/y)", found.rows[0].cells[2].text)

    def test_analysis_summary_table_expands_when_template_has_too_few_rows(self) -> None:
        document = Document()
        table = document.add_table(rows=3, cols=5)
        items = [
            {"check_item": "构件", "position": "M1", "value": "0.80", "case": "OP1", "is_pass": "满足"},
            {"check_item": "节点冲剪（Load）", "position": "J1", "value": "0.90", "case": "EL1", "is_pass": "满足"},
            {"check_item": "节点冲剪（Strength）", "position": "J2", "value": "1.20", "case": "EL2", "is_pass": "不满足"},
        ]

        write_analysis_summary_table(table, items)

        self.assertEqual(4, len(table.rows))
        self.assertEqual("节点冲剪（Strength）", table.rows[3].cells[0].text)
        self.assertEqual("1.20", table.rows[3].cells[2].text)

    def test_chapter4_load_table_formats_large_and_small_numbers(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=9)

        write_basic_case_loads_table(
            table,
            [
                {
                    "label": "CR01",
                    "fx": 12345.678,
                    "fy": 9999.999,
                    "fz": -12.345,
                    "mx": 0,
                    "my": "10000.4",
                    "mz": "abc",
                    "dead_load": "",
                    "buoyancy": -12345.678,
                }
            ],
        )

        row = table.rows[1].cells
        self.assertEqual("12346", row[1].text)
        self.assertEqual("10000.00", row[2].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}12.35", row[3].text)
        self.assertEqual("0.00", row[4].text)
        self.assertEqual("10000", row[5].text)
        self.assertEqual("abc", row[6].text)
        self.assertEqual("", row[7].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}12346", row[8].text)

    def test_combo_case_desc_table_keeps_first_columns_readable(self) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=4)

        write_combo_case_desc_table(
            table,
            [
                {
                    "case": 155,
                    "label": "OP01",
                    "category": "Operation",
                    "desc": "DX00*0.013+DY27*0.000",
                }
            ],
        )

        tbl_pr = table._tbl.tblPr
        self.assertEqual("fixed", tbl_pr.first_child_found_in("w:tblLayout").get(qn("w:type")))
        self.assertFalse(table.autofit)

        expected_widths = ["794", "907", "1191"]
        for index, expected_width in enumerate(expected_widths):
            tc_pr = table.rows[1].cells[index]._tc.get_or_add_tcPr()
            self.assertEqual(expected_width, tc_pr.tcW.get(qn("w:w")))
            self.assertIsNotNone(tc_pr.first_child_found_in("w:noWrap"))

    def test_pile_capacity_table_formats_only_numeric_columns(self) -> None:
        document = Document()
        table = document.add_table(rows=3, cols=10)

        write_pile_capacity_table(
            table,
            [
                {
                    "pile_head_id": "P108",
                    "compression_capacity_kn": "110887.125",
                    "tension_capacity_kn": "9999.1",
                    "pile_weight_kn": "6864.3",
                    "compression_case": "OL12",
                    "compression_load_kn": "38993.3",
                    "tension_case": "OL18",
                    "tension_load_kn": "47123.9",
                    "compression_sf": "2.424",
                    "tension_sf": "2.9",
                }
            ],
        )

        row = table.rows[2].cells
        self.assertEqual("P108", row[0].text)
        self.assertEqual("110887", row[1].text)
        self.assertEqual("9999.10", row[2].text)
        self.assertEqual("6864.30", row[3].text)
        self.assertEqual("OL12", row[4].text)
        self.assertEqual("38993", row[5].text)
        self.assertEqual("OL18", row[6].text)
        self.assertEqual("47124", row[7].text)
        self.assertEqual("2.42", row[8].text)
        self.assertEqual("2.90", row[9].text)

    def test_marine_growth_table_trims_to_actual_layer_count(self) -> None:
        document = Document()
        table = document.add_table(rows=5, cols=11)
        for column_index in range(11):
            table.rows[0].cells[column_index].text = str(column_index)

        write_environment_marine_growth_table(
            table,
            [
                {
                    "layer_no": 1,
                    "upper_limit_m": "0",
                    "lower_limit_m": "-15",
                    "thickness_mm": "10",
                    "density_t_per_m3": "1.4",
                },
                {
                    "layer_no": 4,
                    "upper_limit_m": "-50",
                    "lower_limit_m": "-60",
                    "thickness_mm": "4.5",
                    "density_t_per_m3": "1.4",
                },
            ],
        )

        self.assertEqual(6, len(table.columns))
        self.assertEqual(["0", "1", "1", "2", "3", "4"], [cell.text for cell in table.rows[0].cells])
        self.assertEqual("0", table.rows[1].cells[2].text)
        self.assertEqual("", table.rows[1].cells[3].text)
        self.assertEqual("", table.rows[1].cells[4].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}50", table.rows[1].cells[5].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}60", table.rows[2].cells[5].text)
        self.assertEqual("4.5", table.rows[3].cells[5].text)
        self.assertEqual(["1.4", "1.4", "1.4", "1.4"], [cell.text for cell in table.rows[4].cells[2:]])

    def test_marine_growth_table_expands_beyond_template_layer_count(self) -> None:
        document = Document()
        table = document.add_table(rows=5, cols=11)

        write_environment_marine_growth_table(
            table,
            [
                {
                    "layer_no": 12,
                    "upper_limit_m": "-120",
                    "lower_limit_m": "-130",
                    "thickness_mm": "3",
                    "density_t_per_m3": "1.3",
                }
            ],
        )

        self.assertEqual(14, len(table.columns))
        self.assertEqual("12", table.rows[0].cells[13].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}120", table.rows[1].cells[13].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}130", table.rows[2].cells[13].text)
        self.assertEqual("3", table.rows[3].cells[13].text)
        self.assertEqual("1.3", table.rows[4].cells[13].text)

    def test_marine_growth_table_trims_real_template_with_merged_density_row(self) -> None:
        template_path = next(
            path
            for path in REPORT_MODULE_ROOT.glob("*.docx")
            if len(Document(str(path)).tables) == 14 and "before" not in path.name.lower()
        )
        document = Document(str(template_path))
        table = document.tables[5]

        write_environment_marine_growth_table(
            table,
            [
                {
                    "layer_no": 1,
                    "upper_limit_m": "0",
                    "lower_limit_m": "-15",
                    "thickness_mm": "10",
                    "density_t_per_m3": "1.4",
                },
                {
                    "layer_no": 4,
                    "upper_limit_m": "-50",
                    "lower_limit_m": "-60",
                    "thickness_mm": "4.5",
                    "density_t_per_m3": "1.4",
                },
            ],
        )

        self.assertEqual(6, len(table.columns))
        self.assertEqual("4", table.rows[0].cells[5].text)
        self.assertEqual(f"{NON_BREAKING_HYPHEN}50", table.rows[1].cells[5].text)
        self.assertEqual("1.4", table.rows[4].cells[5].text)


if __name__ == "__main__":
    unittest.main()
