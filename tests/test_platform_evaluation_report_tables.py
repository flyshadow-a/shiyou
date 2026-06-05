from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_MODULE_ROOT = PROJECT_ROOT / "pages" / "output_feasibility_analysis_report"
if str(REPORT_MODULE_ROOT) not in sys.path:
    sys.path.insert(0, str(REPORT_MODULE_ROOT))

from src.renderers.doc_renderer import (  # noqa: E402
    _normalize_chapter_section,
    _render_platform_evaluation_detail_tables,
    _write_platform_evaluation_table_cell,
)


def _first_child(element, tag_name: str):
    return element.find(qn(tag_name))


class PlatformEvaluationReportTablesTests(unittest.TestCase):
    def test_platform_evaluation_rows_survive_chapter_section_normalization(self) -> None:
        result = _normalize_chapter_section(
            {
                "blocks": [{"text": "评估结论"}],
                "well_slot_rows": [{"slot_no": 1, "x": 2}],
                "riser_rows": [{"riser_no": 3, "batter_y": 4}],
                "topside_weight_rows": [{"weight_no": 5, "weight_t": 6}],
            }
        )

        self.assertEqual([{"slot_no": 1, "x": 2}], result["well_slot_rows"])
        self.assertEqual([{"riser_no": 3, "batter_y": 4}], result["riser_rows"])
        self.assertEqual([{"weight_no": 5, "weight_t": 6}], result["topside_weight_rows"])

    def test_platform_evaluation_cells_clear_template_character_indents(self) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        p_pr = paragraph._p.get_or_add_pPr()
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            from docx.oxml import OxmlElement

            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        for attr_name in ("firstLineChars", "leftChars", "rightChars", "hangingChars"):
            ind.set(qn(f"w:{attr_name}"), "200")

        _write_platform_evaluation_table_cell(table.cell(0, 0), "编号")

        ind = paragraph._p.get_or_add_pPr().find(qn("w:ind"))
        self.assertEqual("0", ind.get(qn("w:firstLineChars")))
        self.assertEqual("0", ind.get(qn("w:leftChars")))
        self.assertEqual("0", ind.get(qn("w:rightChars")))
        self.assertEqual("0", ind.get(qn("w:hangingChars")))

    def test_platform_evaluation_detail_tables_use_compact_fixed_layout(self) -> None:
        document = Document()
        anchor = document.add_paragraph(
            "本次改造新增井槽1根，立管和电缆1条，上部组块增加重量12.5t。"
            "对平台改造后的整体结构进行设计水平强度分析，综合以上结果。"
        )

        _render_platform_evaluation_detail_tables(
            document,
            {
                "text": anchor.text,
                "well_slot_rows": [{"slot_no": 1, "x": 2, "top_load_fz": 3}],
                "riser_rows": [{"riser_no": 2, "x": 4, "batter_y": 5}],
                "topside_weight_rows": [{"weight_no": 3, "z": 6, "weight_t": 7}],
            },
        )

        self.assertEqual(3, len(document.tables))
        well_table, riser_table, topside_table = document.tables

        for table in document.tables:
            self.assertFalse(table.autofit)
            tbl_pr = table._tbl.tblPr
            self.assertEqual("0", _first_child(tbl_pr, "w:tblInd").get(qn("w:w")))
            self.assertEqual("fixed", _first_child(tbl_pr, "w:tblLayout").get(qn("w:type")))
            self.assertEqual("8800", _first_child(tbl_pr, "w:tblW").get(qn("w:w")))

            first_cell = table.cell(0, 0)
            tc_mar = first_cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            self.assertEqual("40", tc_mar.find(qn("w:left")).get(qn("w:w")))
            self.assertEqual("40", tc_mar.find(qn("w:right")).get(qn("w:w")))

            run = first_cell.paragraphs[0].runs[0]
            self.assertEqual(Pt(9), run.font.size)

        self.assertEqual(8, len(well_table.columns))
        self.assertEqual(9, len(riser_table.columns))
        self.assertEqual(5, len(topside_table.columns))


if __name__ == "__main__":
    unittest.main()
