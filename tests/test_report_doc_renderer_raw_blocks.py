from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_MODULE_ROOT = PROJECT_ROOT / "pages" / "output_feasibility_analysis_report"
if str(REPORT_MODULE_ROOT) not in sys.path:
    sys.path.insert(0, str(REPORT_MODULE_ROOT))

from src.renderers.doc_renderer import _replace_paragraph_region  # noqa: E402


class ReportDocRendererRawBlockTests(unittest.TestCase):
    def _set_character_indents(self, paragraph) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        ind.set(qn("w:firstLineChars"), "200")
        ind.set(qn("w:leftChars"), "200")
        ind.set(qn("w:hangingChars"), "100")

    def _indent_attrs(self, paragraph) -> dict[str, str | None]:
        ind = paragraph._p.get_or_add_pPr().find(qn("w:ind"))
        return {
            name: ind.get(qn(f"w:{name}")) if ind is not None else None
            for name in ("firstLineChars", "leftChars", "rightChars", "hangingChars")
        }

    def test_replace_paragraph_region_resets_raw_block_indentation(self) -> None:
        document = Document()
        document.add_paragraph("anchor")
        template = document.add_paragraph("template")
        template.paragraph_format.left_indent = Pt(24)
        template.paragraph_format.first_line_indent = Pt(18)
        self._set_character_indents(template)
        document.add_paragraph("next section")

        _replace_paragraph_region(
            document,
            start_index=1,
            end_index=2,
            lines=["4B2 514L-524L OP1D     0.69", "0.85"],
            font_size_pt=7,
        )

        first_raw = document.paragraphs[1]
        second_raw = document.paragraphs[2]
        self.assertEqual("4B2 514L-524L OP1D     0.69", first_raw.text)
        self.assertEqual("0.85", second_raw.text)
        self.assertEqual(0, first_raw.paragraph_format.left_indent.pt)
        self.assertEqual(0, first_raw.paragraph_format.first_line_indent.pt)
        self.assertEqual(0, second_raw.paragraph_format.left_indent.pt)
        self.assertEqual(0, second_raw.paragraph_format.first_line_indent.pt)
        self.assertEqual(
            {"firstLineChars": "0", "leftChars": "0", "rightChars": "0", "hangingChars": "0"},
            self._indent_attrs(first_raw),
        )
        self.assertEqual(
            {"firstLineChars": "0", "leftChars": "0", "rightChars": "0", "hangingChars": "0"},
            self._indent_attrs(second_raw),
        )
        self.assertEqual("Courier New", first_raw.runs[0].font.name)
        self.assertEqual("Courier New", second_raw.runs[0].font.name)


if __name__ == "__main__":
    unittest.main()
